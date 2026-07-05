import asyncio
from base64 import b64encode
from dataclasses import dataclass, field
from datetime import datetime, timedelta
import io
import ipaddress
import logging
import os
import json
import random
import re
import sqlite3
from typing import Any, Literal, Optional
from urllib.parse import urlparse
from zoneinfo import ZoneInfo, ZoneInfoNotFoundError

from bs4 import BeautifulSoup
import discord
from discord.app_commands import Choice
from discord.ext import commands
from discord.ui import LayoutView, TextDisplay
from docx import Document
from dotenv import load_dotenv
import httpx
from openai import AsyncOpenAI
from pypdf import PdfReader
import yaml

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s: %(message)s",
)

VISION_MODEL_TAGS = ("chat-latest", "claude", "gemini", "gemma", "gpt-4", "gpt-5", "gpt-latest", "grok-4", "llama", "vision", "vl")

EMBED_COLOR_COMPLETE = discord.Color.dark_green()
EMBED_COLOR_INCOMPLETE = discord.Color.orange()

STREAMING_INDICATOR = " âšª"
EDIT_DELAY_SECONDS = 1

MAX_MESSAGE_NODES = 500
SETTINGS_FILENAME = "channel_settings.json"
DATABASE_FILENAME = "llmcord.sqlite3"
URL_RE = re.compile(r"https?://[^\s<>()\]\}]+")

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

DEFAULT_MODULES = {
    "server_newspaper": {
        "enabled": False,
        "timezone": "America/New_York",
        "post_time": "00:00",
        "lookback_hours": 24,
        "max_messages_per_channel": 150,
        "max_total_messages": 800,
        "ignored_channel_ids": [],
        "model": None,
    },
    "mediator": {
        "enabled": True,
        "model": None,
        "cooldown_seconds": 60,
    },
    "reputation_titles": {
        "enabled": True,
        "track_messages": True,
        "track_voice": True,
        "track_reactions": True,
        "track_commands": True,
        "ignored_channel_ids": [],
        "spam_cooldown_seconds": 60,
        "allow_user_opt_out": True,
    },
    "party_games": {
        "enabled": True,
        "cooldown_seconds": 20,
        "default_round_timeout_seconds": 60,
        "ai_generated_prompts": True,
        "safe_mode": True,
        "model": None,
        "enabled_games": {
            "would_you_rather": True,
            "never_have_i_ever": True,
            "two_truths_and_lie": True,
            "word_chain": True,
            "trivia": True,
            "guess": True,
            "this_or_that": True,
        },
    },
    "attachment_brain": {
        "enabled": True,
        "max_file_size_mb": 10,
        "max_extracted_chars": 50000,
        "chunk_size_chars": 12000,
        "store_file_contents": False,
        "allowed_extensions": [
            "txt", "md", "pdf", "docx", "csv", "json", "log", "py", "js", "ts", "html", "css", "lua",
            "java", "cs", "cpp", "c", "go", "rs", "yaml", "yml", "toml", "ini", "png", "jpg", "jpeg", "webp",
        ],
        "vision_model": None,
        "model": None,
    },
}

SERIOUS_MEDIATION_RE = re.compile(
    r"\b(abuse|assault|blackmail|custody|divorce|emergency|eviction|harassment|hospital|illegal|lawsuit|legal|"
    r"medical|mental health|police|restraining order|self[- ]?harm|suicide|threat|violence|weapon)\b",
    re.IGNORECASE,
)
REFUSE_MEDIATION_RE = re.compile(
    r"\b(abuse|assault|emergency|harassment|kill|self[- ]?harm|suicide|threat|violence|weapon)\b",
    re.IGNORECASE,
)
def parse_env_value(value: Optional[str]) -> Any:
    if value is None:
        return None

    stripped = value.strip()

    # Allow Render env vars like ADMIN_IDS=[123,456] or ADMIN_IDS=123,456
    try:
        return json.loads(stripped)
    except json.JSONDecodeError:
        pass

    if "," in stripped:
        parts = [part.strip() for part in stripped.split(",") if part.strip()]
        return [int(part) if part.isdigit() else part for part in parts]

    return int(stripped) if stripped.isdigit() else value


def resolve_env(node: Any) -> Any:
    if isinstance(node, dict):
        return {key.removesuffix("_env"): parse_env_value(os.environ.get(value)) if key.endswith("_env") else resolve_env(value) for key, value in node.items()}
    return node


def get_config(filename: str = "config.yaml") -> dict[str, Any]:
    with open(filename, encoding="utf-8") as file:
        return resolve_env(yaml.safe_load(file))


def get_db() -> sqlite3.Connection:
    conn = sqlite3.connect(DATABASE_FILENAME)
    conn.row_factory = sqlite3.Row
    return conn


def init_db() -> None:
    with get_db() as conn:
        conn.executescript(
            """
            CREATE TABLE IF NOT EXISTS module_settings (
                guild_id INTEGER NOT NULL,
                module TEXT NOT NULL,
                key TEXT NOT NULL,
                value TEXT NOT NULL,
                PRIMARY KEY (guild_id, module, key)
            );

            CREATE TABLE IF NOT EXISTS module_ignored_channels (
                guild_id INTEGER NOT NULL,
                module TEXT NOT NULL,
                channel_id INTEGER NOT NULL,
                PRIMARY KEY (guild_id, module, channel_id)
            );

            CREATE TABLE IF NOT EXISTS newspaper_runs (
                guild_id INTEGER NOT NULL,
                issue_date TEXT NOT NULL,
                posted_at TEXT NOT NULL,
                channel_id INTEGER,
                PRIMARY KEY (guild_id, issue_date)
            );

            CREATE TABLE IF NOT EXISTS reputation_title_stats (
                guild_id INTEGER NOT NULL,
                user_id INTEGER NOT NULL,
                stat TEXT NOT NULL,
                value INTEGER NOT NULL DEFAULT 0,
                updated_at TEXT NOT NULL,
                PRIMARY KEY (guild_id, user_id, stat)
            );

            CREATE TABLE IF NOT EXISTS reputation_title_user_titles (
                guild_id INTEGER NOT NULL,
                user_id INTEGER NOT NULL,
                title TEXT NOT NULL,
                earned_at TEXT NOT NULL,
                granted_by INTEGER,
                source TEXT NOT NULL DEFAULT 'auto',
                PRIMARY KEY (guild_id, user_id, title)
            );

            CREATE TABLE IF NOT EXISTS reputation_title_profiles (
                guild_id INTEGER NOT NULL,
                user_id INTEGER NOT NULL,
                equipped_title TEXT,
                opted_out_at TEXT,
                PRIMARY KEY (guild_id, user_id)
            );

            CREATE TABLE IF NOT EXISTS party_game_scores (
                guild_id INTEGER NOT NULL,
                user_id INTEGER NOT NULL,
                game TEXT NOT NULL,
                score INTEGER NOT NULL DEFAULT 0,
                streak INTEGER NOT NULL DEFAULT 0,
                updated_at TEXT NOT NULL,
                PRIMARY KEY (guild_id, user_id, game)
            );

            CREATE TABLE IF NOT EXISTS party_game_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                guild_id INTEGER NOT NULL,
                channel_id INTEGER NOT NULL,
                game TEXT NOT NULL,
                user_id INTEGER,
                result TEXT,
                created_at TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS party_game_states (
                guild_id INTEGER NOT NULL,
                channel_id INTEGER NOT NULL,
                game TEXT NOT NULL,
                state TEXT NOT NULL,
                expires_at TEXT,
                created_at TEXT NOT NULL,
                PRIMARY KEY (guild_id, channel_id, game)
            );
            """
        )


def json_dumps(value: Any) -> str:
    return json.dumps(value, separators=(",", ":"))


def json_loads(value: Optional[str], default: Any = None) -> Any:
    if value is None:
        return default

    try:
        return json.loads(value)
    except json.JSONDecodeError:
        return default


def now_iso() -> str:
    return datetime.now().astimezone().isoformat()


def get_module_config(loaded_config: dict[str, Any], module: str, guild_id: Optional[int] = None) -> dict[str, Any]:
    module_config = dict(DEFAULT_MODULES.get(module, {}))
    module_config.update((loaded_config.get("modules") or {}).get(module) or {})

    if guild_id is not None:
        with get_db() as conn:
            rows = conn.execute(
                "SELECT key, value FROM module_settings WHERE guild_id = ? AND module = ?",
                (guild_id, module),
            ).fetchall()
            for row in rows:
                module_config[row["key"]] = json_loads(row["value"], row["value"])

            ignored_rows = conn.execute(
                "SELECT channel_id FROM module_ignored_channels WHERE guild_id = ? AND module = ?",
                (guild_id, module),
            ).fetchall()

        ignored_channel_ids = set(int(channel_id) for channel_id in module_config.get("ignored_channel_ids", []) or [])
        ignored_channel_ids.update(int(row["channel_id"]) for row in ignored_rows)
        module_config["ignored_channel_ids"] = sorted(ignored_channel_ids)

    return module_config


def set_module_setting(guild_id: int, module: str, key: str, value: Any) -> None:
    with get_db() as conn:
        conn.execute(
            "INSERT OR REPLACE INTO module_settings (guild_id, module, key, value) VALUES (?, ?, ?, ?)",
            (guild_id, module, key, json_dumps(value)),
        )


def get_module_setting(guild_id: int, module: str, key: str, default: Any = None) -> Any:
    with get_db() as conn:
        row = conn.execute(
            "SELECT value FROM module_settings WHERE guild_id = ? AND module = ? AND key = ?",
            (guild_id, module, key),
        ).fetchone()
    return json_loads(row["value"], default) if row else default


def set_ignored_channel(guild_id: int, module: str, channel_id: int, ignored: bool) -> None:
    with get_db() as conn:
        if ignored:
            conn.execute(
                "INSERT OR IGNORE INTO module_ignored_channels (guild_id, module, channel_id) VALUES (?, ?, ?)",
                (guild_id, module, channel_id),
            )
        else:
            conn.execute(
                "DELETE FROM module_ignored_channels WHERE guild_id = ? AND module = ? AND channel_id = ?",
                (guild_id, module, channel_id),
            )


def load_channel_settings(filename: str = SETTINGS_FILENAME) -> dict[str, Any]:
    if not os.path.exists(filename):
        return {"channels": {}}

    with open(filename, encoding="utf-8") as file:
        loaded_settings = json.load(file)

    loaded_settings.setdefault("channels", {})
    return loaded_settings


def save_channel_settings(filename: str = SETTINGS_FILENAME) -> None:
    with open(filename, "w", encoding="utf-8") as file:
        json.dump(channel_settings, file, indent=2, sort_keys=True)


def get_channel_key(channel: Any) -> Optional[str]:
    channel_id = getattr(channel, "id", None)
    return str(channel_id) if channel_id is not None else None


def get_channel_override(channel: Any, key: str) -> Optional[str]:
    channel_key = get_channel_key(channel)
    if channel_key is None:
        return None
    value = channel_settings.get("channels", {}).get(channel_key, {}).get(key)
    return value if isinstance(value, str) else None


def set_channel_override(channel: Any, key: str, value: Optional[str]) -> None:
    channel_key = get_channel_key(channel)
    if channel_key is None:
        return

    channel_values = channel_settings.setdefault("channels", {}).setdefault(channel_key, {})
    if value is None:
        channel_values.pop(key, None)
    else:
        channel_values[key] = value

    if not channel_values:
        channel_settings["channels"].pop(channel_key, None)

    save_channel_settings()


def get_effective_model(channel: Any, loaded_config: dict[str, Any]) -> str:
    model = get_channel_override(channel, "model")
    return model if model in loaded_config["models"] else curr_model


def get_personas(loaded_config: dict[str, Any]) -> dict[str, str]:
    personas = {"default": loaded_config.get("system_prompt", "") or ""}
    personas.update(loaded_config.get("personas") or {})
    return {name: prompt for name, prompt in personas.items() if isinstance(prompt, str)}


def get_effective_persona(channel: Any, loaded_config: dict[str, Any]) -> str:
    persona = get_channel_override(channel, "persona")
    return persona if persona in get_personas(loaded_config) else "default"


config = get_config()
init_db()
curr_model = next(iter(config["models"]))
channel_settings = load_channel_settings()

msg_nodes = {}
mediator_cooldowns = {}
reputation_cooldowns = {}
party_cooldowns = {}
party_rounds = {}
newspaper_task: Optional[asyncio.Task] = None
last_task_time = 0

intents = discord.Intents.default()
intents.message_content = True
activity = discord.CustomActivity(name=(config.get("status_message") or "github.com/jakobdylanc/llmcord")[:128])
discord_bot = commands.Bot(intents=intents, activity=activity, command_prefix=None)

httpx_client = httpx.AsyncClient()


@dataclass
class MsgNode:
    role: Literal["user", "assistant"] = "assistant"

    text: Optional[str] = None
    images: list[dict[str, Any]] = field(default_factory=list)

    has_bad_attachments: bool = False
    has_bad_links: bool = False
    fetch_parent_failed: bool = False

    parent_msg: Optional[discord.Message] = None

    lock: asyncio.Lock = field(default_factory=asyncio.Lock)



def interaction_is_private(interaction: discord.Interaction) -> bool:
    return interaction.channel is None or getattr(interaction.channel, "type", None) == discord.ChannelType.private


def is_admin_user(user_id: int, loaded_config: dict[str, Any]) -> bool:
    admin_ids = loaded_config.get("permissions", {}).get("users", {}).get("admin_ids", []) or []
    if isinstance(admin_ids, int):
        admin_ids = [admin_ids]
    return user_id in {int(admin_id) for admin_id in admin_ids}


def user_has_permission_for_interaction(interaction: discord.Interaction, loaded_config: dict[str, Any]) -> bool:
    permissions = loaded_config["permissions"]
    user_id = interaction.user.id

    if is_admin_user(user_id, loaded_config):
        return True

    allowed_user_ids = permissions["users"].get("allowed_ids", []) or []
    blocked_user_ids = permissions["users"].get("blocked_ids", []) or []
    allowed_role_ids = permissions["roles"].get("allowed_ids", []) or []
    blocked_role_ids = permissions["roles"].get("blocked_ids", []) or []
    allowed_channel_ids = permissions["channels"].get("allowed_ids", []) or []
    blocked_channel_ids = permissions["channels"].get("blocked_ids", []) or []

    role_ids = {role.id for role in getattr(interaction.user, "roles", ())}
    channel = interaction.channel
    channel_ids = set(filter(None, (
        getattr(channel, "id", None),
        getattr(channel, "parent_id", None),
        getattr(channel, "category_id", None),
    )))

    is_dm_or_group = interaction.guild is None
    allow_dms = loaded_config.get("allow_dms", True)

    allow_all_users = not allowed_user_ids if is_dm_or_group else not allowed_user_ids and not allowed_role_ids
    is_good_user = allow_all_users or user_id in allowed_user_ids or any(role_id in allowed_role_ids for role_id in role_ids)
    is_bad_user = not is_good_user or user_id in blocked_user_ids or any(role_id in blocked_role_ids for role_id in role_ids)

    allow_all_channels = not allowed_channel_ids
    is_good_channel = allow_dms if is_dm_or_group else allow_all_channels or any(channel_id in allowed_channel_ids for channel_id in channel_ids)
    is_bad_channel = not is_good_channel or any(channel_id in blocked_channel_ids for channel_id in channel_ids)

    return not is_bad_user and not is_bad_channel


def get_attachment_kind(attachment: discord.Attachment) -> Optional[Literal["text", "image", "pdf", "docx"]]:
    content_type = (attachment.content_type or "").lower()
    filename = attachment.filename.lower()

    if content_type.startswith("text") or filename.endswith((".txt", ".md", ".py", ".js", ".ts", ".tsx", ".jsx", ".json", ".yaml", ".yml", ".csv", ".log")):
        return "text"
    if content_type.startswith("image"):
        return "image"
    if content_type == "application/pdf" or filename.endswith(".pdf"):
        return "pdf"
    if content_type == DOCX_CONTENT_TYPE or filename.endswith(".docx"):
        return "docx"
    return None


def normalize_extracted_text(text: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", text.replace("\r\n", "\n").replace("\r", "\n")).strip()


def extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    return normalize_extracted_text("\n\n".join(page.extract_text() or "" for page in reader.pages))


def extract_docx_text(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return normalize_extracted_text("\n".join(parts))


def extract_html_text(data: bytes) -> str:
    soup = BeautifulSoup(data, "html.parser")

    for tag in soup(["script", "style", "noscript", "svg"]):
        tag.decompose()

    title = soup.title.string.strip() if soup.title and soup.title.string else ""
    body = normalize_extracted_text(soup.get_text("\n"))
    return normalize_extracted_text("\n\n".join(part for part in (title, body) if part))


def extract_response_text(url: str, content_type: str, data: bytes) -> str:
    lowered_url = url.lower().split("?", 1)[0]
    lowered_type = content_type.lower()

    if "application/pdf" in lowered_type or lowered_url.endswith(".pdf"):
        return extract_pdf_text(data)
    if DOCX_CONTENT_TYPE in lowered_type or lowered_url.endswith(".docx"):
        return extract_docx_text(data)
    if "html" in lowered_type:
        return extract_html_text(data)
    if lowered_type.startswith("text/") or "json" in lowered_type or "xml" in lowered_type:
        return normalize_extracted_text(data.decode("utf-8", errors="replace"))
    return ""


def is_fetchable_url(url: str) -> bool:
    parsed_url = urlparse(url)
    if parsed_url.scheme not in ("http", "https") or not parsed_url.hostname:
        return False

    hostname = parsed_url.hostname.lower()
    if hostname in ("localhost", "127.0.0.1", "::1") or hostname.endswith(".local"):
        return False

    try:
        ip_address = ipaddress.ip_address(hostname)
    except ValueError:
        return True

    return not (ip_address.is_private or ip_address.is_loopback or ip_address.is_link_local or ip_address.is_multicast)


async def extract_url_texts(text: str, loaded_config: dict[str, Any]) -> tuple[list[str], bool]:
    max_urls = loaded_config.get("max_urls", 3)
    max_url_text = loaded_config.get("max_url_text", 15000)
    timeout = loaded_config.get("url_fetch_timeout", 10)
    urls = []
    had_failures = False

    for match in URL_RE.finditer(text):
        url = match.group(0).rstrip(".,;:!?\"'")
        if not is_fetchable_url(url):
            had_failures = True
            continue
        if url not in urls:
            urls.append(url)
        if len(urls) >= max_urls:
            break

    url_texts = []

    for url in urls:
        try:
            response = await httpx_client.get(
                url,
                follow_redirects=True,
                timeout=timeout,
                headers={"User-Agent": "llmcord/1.0"},
            )
            response.raise_for_status()

            extracted_text = await asyncio.to_thread(
                extract_response_text,
                str(response.url),
                response.headers.get("content-type", ""),
                response.content,
            )
            if extracted_text:
                url_texts.append(f"[URL: {url}]\n{extracted_text[:max_url_text]}")
            else:
                had_failures = True
        except Exception:
            logging.exception("Error fetching URL for message context: %s", url)
            had_failures = True

    return url_texts, had_failures


async def extract_attachment_text(attachment: discord.Attachment, response: httpx.Response, kind: str, loaded_config: dict[str, Any]) -> str:
    max_attachment_text = loaded_config.get("max_attachment_text", loaded_config.get("max_text", 100000))

    if kind == "text":
        return response.text[:max_attachment_text]
    if kind == "pdf":
        text = await asyncio.to_thread(extract_pdf_text, response.content)
        return f"[PDF: {attachment.filename}]\n{text[:max_attachment_text]}"
    if kind == "docx":
        text = await asyncio.to_thread(extract_docx_text, response.content)
        return f"[DOCX: {attachment.filename}]\n{text[:max_attachment_text]}"
    return ""


async def populate_msg_node(curr_msg: discord.Message, curr_node: MsgNode, loaded_config: dict[str, Any]) -> None:
    cleaned_content = curr_msg.content.removeprefix(discord_bot.user.mention).lstrip()
    curr_node.role = "assistant" if curr_msg.author == discord_bot.user else "user"

    attachment_kinds = [(att, kind) for att in curr_msg.attachments if (kind := get_attachment_kind(att))]
    attachment_responses = await asyncio.gather(
        *[httpx_client.get(att.url) for att, _ in attachment_kinds],
        return_exceptions=True,
    )

    attachment_texts = []
    curr_node.images = []
    bad_attachment_count = len(curr_msg.attachments) - len(attachment_kinds)

    for (attachment, kind), response in zip(attachment_kinds, attachment_responses):
        if isinstance(response, Exception):
            logging.warning("Error fetching attachment for message context: %s", response)
            bad_attachment_count += 1
            continue

        try:
            if kind == "image":
                curr_node.images.append(dict(type="image_url", image_url=dict(url=f"data:{attachment.content_type};base64,{b64encode(response.content).decode('utf-8')}")))
            else:
                attachment_text = await extract_attachment_text(attachment, response, kind, loaded_config)
                if attachment_text:
                    attachment_texts.append(attachment_text)
        except Exception:
            logging.exception("Error extracting attachment text")
            bad_attachment_count += 1

    url_texts = []
    if curr_node.role == "user" and cleaned_content:
        url_texts, curr_node.has_bad_links = await extract_url_texts(cleaned_content, loaded_config)

    curr_node.text = "\n".join(
        ([cleaned_content] if cleaned_content else [])
        + ["\n".join(filter(None, (embed.title, embed.description, embed.footer.text))) for embed in curr_msg.embeds]
        + [component.content for component in curr_msg.components if component.type == discord.ComponentType.text_display]
        + attachment_texts
        + url_texts
    )

    if curr_node.role == "user" and (curr_node.text or curr_node.images):
        curr_node.text = f"<@{curr_msg.author.id}>: {curr_node.text}"

    curr_node.has_bad_attachments = bad_attachment_count > 0


def node_to_message_content(curr_node: MsgNode, max_text: int, max_images: int) -> Any:
    text = (curr_node.text or "")[:max_text]

    if curr_node.images[:max_images]:
        return [dict(type="text", text=text)] + curr_node.images[:max_images]
    return text


def build_openai_client_and_kwargs(loaded_config: dict[str, Any], provider_slash_model: str, messages: list[dict[str, Any]], stream: bool) -> dict[str, Any]:
    provider, model = provider_slash_model.removesuffix(":vision").split("/", 1)

    provider_config = loaded_config["providers"][provider]
    base_url = provider_config["base_url"]
    api_key = provider_config.get("api_key", "sk-no-key-required")
    openai_client = AsyncOpenAI(base_url=base_url, api_key=api_key)

    model_parameters = loaded_config["models"].get(provider_slash_model, None)
    extra_headers = provider_config.get("extra_headers")
    extra_query = provider_config.get("extra_query")
    extra_body = (provider_config.get("extra_body") or {}) | (model_parameters or {}) or None

    return dict(
        client=openai_client,
        kwargs=dict(
            model=model,
            messages=messages,
            stream=stream,
            extra_headers=extra_headers,
            extra_query=extra_query,
            extra_body=extra_body,
        ),
    )


def get_system_prompt(loaded_config: dict[str, Any], persona: str = "default") -> str:
    system_prompt = get_personas(loaded_config).get(persona) or ""
    if system_prompt:
        now = datetime.now().astimezone()
        system_prompt = system_prompt.replace("{date}", now.strftime("%B %d %Y")).replace("{time}", now.strftime("%H:%M:%S %Z%z")).strip()
    return system_prompt


def append_system_prompt(messages: list[dict[str, Any]], loaded_config: dict[str, Any], persona: str = "default") -> None:
    if system_prompt := get_system_prompt(loaded_config, persona):
        messages.append(dict(role="system", content=system_prompt))


async def generate_nonstream_response(prompt: str, user_id: int, loaded_config: dict[str, Any], provider_slash_model: str, persona: str) -> str:
    messages = [dict(role="user", content=f"<@{user_id}>: {prompt[:loaded_config.get('max_text', 100000)]}")]
    append_system_prompt(messages, loaded_config, persona)

    client_and_kwargs = build_openai_client_and_kwargs(loaded_config, provider_slash_model, messages[::-1], stream=False)
    response = await client_and_kwargs["client"].chat.completions.create(**client_and_kwargs["kwargs"])
    return response.choices[0].message.content or ""


async def send_interaction_chunks(interaction: discord.Interaction, content: str, private: bool) -> None:
    max_len = 1900
    chunks = [content[i:i + max_len] for i in range(0, len(content), max_len)] or ["*(empty response)*"]
    for index, chunk in enumerate(chunks):
        if index == 0:
            await interaction.followup.send(chunk, ephemeral=private)
        else:
            await interaction.followup.send(chunk, ephemeral=private)


async def set_parent_msg(curr_msg: discord.Message, curr_node: MsgNode) -> None:
    try:
        if (
            curr_msg.reference == None
            and discord_bot.user.mention not in curr_msg.content
            and (prev_msg_in_channel := ([m async for m in curr_msg.channel.history(before=curr_msg, limit=1)] or [None])[0])
            and prev_msg_in_channel.type in (discord.MessageType.default, discord.MessageType.reply)
            and prev_msg_in_channel.author == (discord_bot.user if curr_msg.channel.type == discord.ChannelType.private else curr_msg.author)
        ):
            curr_node.parent_msg = prev_msg_in_channel
        else:
            is_public_thread = curr_msg.channel.type == discord.ChannelType.public_thread
            parent_is_thread_start = is_public_thread and curr_msg.reference == None and curr_msg.channel.parent.type == discord.ChannelType.text

            if parent_msg_id := curr_msg.channel.id if parent_is_thread_start else getattr(curr_msg.reference, "message_id", None):
                if parent_is_thread_start:
                    curr_node.parent_msg = curr_msg.channel.starter_message or await curr_msg.channel.parent.fetch_message(parent_msg_id)
                else:
                    curr_node.parent_msg = curr_msg.reference.cached_message or await curr_msg.channel.fetch_message(parent_msg_id)

    except (discord.NotFound, discord.HTTPException):
        logging.exception("Error fetching next message in the chain")
        curr_node.fetch_parent_failed = True


async def build_reply_chain_messages(start_msg: discord.Message, loaded_config: dict[str, Any], accept_images: bool) -> tuple[list[dict[str, Any]], set[str]]:
    max_text = loaded_config.get("max_text", 100000)
    max_images = loaded_config.get("max_images", 5) if accept_images else 0
    max_messages = loaded_config.get("max_messages", 25)

    messages = []
    user_warnings = set()
    curr_msg = start_msg

    while curr_msg != None and len(messages) < max_messages:
        curr_node = msg_nodes.setdefault(curr_msg.id, MsgNode())

        async with curr_node.lock:
            if curr_node.text == None:
                await populate_msg_node(curr_msg, curr_node, loaded_config)
                await set_parent_msg(curr_msg, curr_node)

            content = node_to_message_content(curr_node, max_text, max_images)

            if content != "":
                messages.append(dict(content=content, role=curr_node.role))

            if len(curr_node.text or "") > max_text:
                user_warnings.add(f"Warning: Max {max_text:,} characters per message")
            if len(curr_node.images) > max_images:
                user_warnings.add(f"Warning: Max {max_images} image{'' if max_images == 1 else 's'} per message" if max_images > 0 else "Warning: Can't see images")
            if curr_node.has_bad_attachments:
                user_warnings.add("Warning: Unsupported or unreadable attachments")
            if curr_node.has_bad_links:
                user_warnings.add("Warning: Some URLs could not be read")
            if curr_node.fetch_parent_failed or (curr_node.parent_msg != None and len(messages) == max_messages):
                user_warnings.add(f"Warning: Only using last {len(messages)} message{'' if len(messages) == 1 else 's'}")

            curr_msg = curr_node.parent_msg

    return messages, user_warnings


async def send_streaming_reply(start_msg: discord.Message, loaded_config: dict[str, Any], log_label: str = "Message received") -> None:
    global last_task_time

    provider_slash_model = get_effective_model(start_msg.channel, loaded_config)
    provider, model = provider_slash_model.removesuffix(":vision").split("/", 1)

    provider_config = loaded_config["providers"][provider]
    base_url = provider_config["base_url"]
    api_key = provider_config.get("api_key", "sk-no-key-required")
    openai_client = AsyncOpenAI(base_url=base_url, api_key=api_key)

    model_parameters = loaded_config["models"].get(provider_slash_model, None)
    extra_headers = provider_config.get("extra_headers")
    extra_query = provider_config.get("extra_query")
    extra_body = (provider_config.get("extra_body") or {}) | (model_parameters or {}) or None

    accept_images = any(x in provider_slash_model.lower() for x in VISION_MODEL_TAGS)
    messages, user_warnings = await build_reply_chain_messages(start_msg, loaded_config, accept_images)

    logging.info(f"{log_label} (user ID: {start_msg.author.id}, attachments: {len(start_msg.attachments)}, conversation length: {len(messages)}, model: {provider_slash_model}):\n{start_msg.content}")

    append_system_prompt(messages, loaded_config, get_effective_persona(start_msg.channel, loaded_config))

    curr_content = finish_reason = None
    response_msgs = []
    response_contents = []

    openai_kwargs = dict(model=model, messages=messages[::-1], stream=True, extra_headers=extra_headers, extra_query=extra_query, extra_body=extra_body)

    if use_plain_responses := loaded_config.get("use_plain_responses", False):
        max_message_length = 4000
    else:
        max_message_length = 4096 - len(STREAMING_INDICATOR)
        embed = discord.Embed.from_dict(dict(fields=[dict(name=warning, value="", inline=False) for warning in sorted(user_warnings)]))

    async def reply_helper(**reply_kwargs) -> None:
        reply_target = start_msg if not response_msgs else response_msgs[-1]
        response_msg = await reply_target.reply(**reply_kwargs)
        response_msgs.append(response_msg)

        msg_nodes[response_msg.id] = MsgNode(parent_msg=start_msg)
        await msg_nodes[response_msg.id].lock.acquire()

    try:
        async with start_msg.channel.typing():
            async for chunk in await openai_client.chat.completions.create(**openai_kwargs):
                if finish_reason != None:
                    break

                if not (choice := chunk.choices[0] if chunk.choices else None):
                    continue

                finish_reason = choice.finish_reason

                prev_content = curr_content or ""
                curr_content = choice.delta.content or ""

                new_content = prev_content if finish_reason == None else (prev_content + curr_content)

                if response_contents == [] and new_content == "":
                    continue

                if start_next_msg := response_contents == [] or len(response_contents[-1] + new_content) > max_message_length:
                    response_contents.append("")

                response_contents[-1] += new_content

                if not use_plain_responses:
                    time_delta = datetime.now().timestamp() - last_task_time

                    ready_to_edit = time_delta >= EDIT_DELAY_SECONDS
                    msg_split_incoming = finish_reason == None and len(response_contents[-1] + curr_content) > max_message_length
                    is_final_edit = finish_reason != None or msg_split_incoming
                    is_good_finish = finish_reason != None and finish_reason.lower() in ("stop", "end_turn")

                    if start_next_msg or ready_to_edit or is_final_edit:
                        embed.description = response_contents[-1] if is_final_edit else (response_contents[-1] + STREAMING_INDICATOR)
                        embed.color = EMBED_COLOR_COMPLETE if msg_split_incoming or is_good_finish else EMBED_COLOR_INCOMPLETE

                        if start_next_msg:
                            await reply_helper(embed=embed, silent=True)
                        else:
                            await asyncio.sleep(EDIT_DELAY_SECONDS - time_delta)
                            await response_msgs[-1].edit(embed=embed)

                        last_task_time = datetime.now().timestamp()

            if use_plain_responses:
                for content in response_contents:
                    await reply_helper(view=LayoutView().add_item(TextDisplay(content=content)))

    except Exception:
        logging.exception("Error while generating response")

    for response_msg in response_msgs:
        msg_nodes[response_msg.id].text = "".join(response_contents)
        msg_nodes[response_msg.id].lock.release()

    if (num_nodes := len(msg_nodes)) > MAX_MESSAGE_NODES:
        for msg_id in sorted(msg_nodes.keys())[: num_nodes - MAX_MESSAGE_NODES]:
            async with msg_nodes.setdefault(msg_id, MsgNode()).lock:
                msg_nodes.pop(msg_id, None)


async def get_referenced_message(message: discord.Message) -> Optional[discord.Message]:
    if node := msg_nodes.get(message.id):
        if node.parent_msg is not None:
            return node.parent_msg

    if not message.reference or not getattr(message.reference, "message_id", None):
        return None

    try:
        return message.reference.cached_message or await message.channel.fetch_message(message.reference.message_id)
    except (discord.NotFound, discord.HTTPException):
        return None


async def find_retry_target(channel: Any) -> Optional[discord.Message]:
    async for message in channel.history(limit=50):
        if message.author != discord_bot.user:
            continue

        seen_msg_ids = set()
        curr_msg = message

        while curr_msg and curr_msg.id not in seen_msg_ids:
            seen_msg_ids.add(curr_msg.id)
            parent_msg = await get_referenced_message(curr_msg)
            if parent_msg is None:
                break
            if parent_msg.author != discord_bot.user:
                return parent_msg
            curr_msg = parent_msg

    return None


async def build_recent_channel_messages(channel: Any, loaded_config: dict[str, Any], limit: int, accept_images: bool) -> tuple[list[dict[str, Any]], set[str]]:
    max_text = loaded_config.get("max_text", 100000)
    max_images = loaded_config.get("max_images", 5) if accept_images else 0
    messages = []
    user_warnings = set()

    async for curr_msg in channel.history(limit=limit):
        if curr_msg.type not in (discord.MessageType.default, discord.MessageType.reply):
            continue

        curr_node = msg_nodes.setdefault(curr_msg.id, MsgNode())

        async with curr_node.lock:
            if curr_node.text == None:
                await populate_msg_node(curr_msg, curr_node, loaded_config)

            content = node_to_message_content(curr_node, max_text, max_images)
            if content != "":
                messages.append(dict(content=content, role=curr_node.role))

            if curr_node.has_bad_attachments:
                user_warnings.add("Warning: Some attachments could not be read")
            if curr_node.has_bad_links:
                user_warnings.add("Warning: Some URLs could not be read")

    return messages, user_warnings


def require_guild(interaction: discord.Interaction) -> Optional[discord.Guild]:
    return interaction.guild


async def require_admin_interaction(interaction: discord.Interaction, loaded_config: dict[str, Any]) -> bool:
    if interaction.guild is None:
        await interaction.response.send_message("This command only works in a server.", ephemeral=True)
        return False
    if not is_admin_user(interaction.user.id, loaded_config):
        await interaction.response.send_message("You don't have permission to do that.", ephemeral=True)
        return False
    return True


def module_model(loaded_config: dict[str, Any], module_config: dict[str, Any], channel: Any) -> str:
    configured_model = module_config.get("model")
    return configured_model if configured_model in loaded_config["models"] else get_effective_model(channel, loaded_config)


async def generate_module_text(loaded_config: dict[str, Any], provider_slash_model: str, messages: list[dict[str, Any]]) -> str:
    client_and_kwargs = build_openai_client_and_kwargs(loaded_config, provider_slash_model, messages, stream=False)
    response = await client_and_kwargs["client"].chat.completions.create(**client_and_kwargs["kwargs"])
    return response.choices[0].message.content or ""


async def send_channel_chunks(channel: Any, content: str) -> list[discord.Message]:
    chunks = [content[i:i + 1900] for i in range(0, len(content), 1900)] or ["*(empty response)*"]
    sent_messages = []
    for chunk in chunks:
        sent_messages.append(await channel.send(chunk))
    return sent_messages


def is_public_server_text_channel(channel: Any, ignored_channel_ids: set[int]) -> bool:
    if not isinstance(channel, (discord.TextChannel, discord.VoiceChannel, discord.StageChannel)):
        return False
    if not isinstance(channel, discord.TextChannel):
        return False
    if channel.id in ignored_channel_ids:
        return False

    guild = channel.guild
    bot_member = guild.me
    if bot_member is None:
        return False

    bot_perms = channel.permissions_for(bot_member)
    everyone_perms = channel.permissions_for(guild.default_role)

    return bool(
        bot_perms.view_channel
        and bot_perms.read_message_history
        and everyone_perms.view_channel
    )


def safe_message_text(message: discord.Message, max_len: int = 500) -> str:
    content = normalize_extracted_text(message.content)
    content = URL_RE.sub("[link]", content)
    content = re.sub(r"<@!?\d+>", "@user", content)
    content = re.sub(r"<#\d+>", "#channel", content)
    return content[:max_len]


async def collect_public_activity(
    guild: discord.Guild,
    module_config: dict[str, Any],
    since: datetime,
    max_total_messages: int,
) -> tuple[list[str], dict[str, int]]:
    ignored_channel_ids = set(int(channel_id) for channel_id in module_config.get("ignored_channel_ids", []) or [])
    max_per_channel = int(module_config.get("max_messages_per_channel", 150))
    activity_lines = []
    channel_counts = {}

    for channel in guild.text_channels:
        if not is_public_server_text_channel(channel, ignored_channel_ids):
            continue

        channel_count = 0
        try:
            async for message in channel.history(after=since, limit=max_per_channel, oldest_first=True):
                if message.author.bot or not message.content:
                    continue
                text = safe_message_text(message)
                if not text:
                    continue

                activity_lines.append(f"#{channel.name} | {message.author.display_name}: {text}")
                channel_count += 1

                if len(activity_lines) >= max_total_messages:
                    break
        except (discord.Forbidden, discord.HTTPException):
            logging.exception("Skipping channel during activity collection: %s", channel.id)

        if channel_count:
            channel_counts[channel.name] = channel_count

        if len(activity_lines) >= max_total_messages:
            break

    return activity_lines, channel_counts


def newspaper_issue_date(timezone_name: str) -> str:
    try:
        tzinfo = ZoneInfo(timezone_name)
    except ZoneInfoNotFoundError:
        tzinfo = ZoneInfo("America/New_York")
    return datetime.now(tzinfo).date().isoformat()


def newspaper_already_posted(guild_id: int, issue_date: str) -> bool:
    with get_db() as conn:
        row = conn.execute(
            "SELECT 1 FROM newspaper_runs WHERE guild_id = ? AND issue_date = ?",
            (guild_id, issue_date),
        ).fetchone()
    return row is not None


def record_newspaper_run(guild_id: int, issue_date: str, channel_id: int) -> None:
    with get_db() as conn:
        conn.execute(
            "INSERT OR REPLACE INTO newspaper_runs (guild_id, issue_date, posted_at, channel_id) VALUES (?, ?, ?, ?)",
            (guild_id, issue_date, now_iso(), channel_id),
        )


async def generate_newspaper(guild: discord.Guild, channel: discord.TextChannel, loaded_config: dict[str, Any]) -> str:
    module_config = get_module_config(loaded_config, "server_newspaper", guild.id)
    timezone_name = module_config.get("timezone", "America/New_York")

    try:
        tzinfo = ZoneInfo(timezone_name)
    except ZoneInfoNotFoundError:
        tzinfo = ZoneInfo("America/New_York")
        timezone_name = "America/New_York"

    lookback_hours = int(module_config.get("lookback_hours", 24))
    since = datetime.now(tzinfo).astimezone() - timedelta(hours=lookback_hours)
    activity_lines, channel_counts = await collect_public_activity(
        guild,
        module_config,
        since,
        int(module_config.get("max_total_messages", 800)),
    )

    active_channels = ", ".join(f"#{name} ({count})" for name, count in sorted(channel_counts.items(), key=lambda item: item[1], reverse=True)[:8])
    if not activity_lines:
        activity_lines = ["Quiet day: not enough public, readable messages were available for a full edition."]

    prompt = (
        "Write a fun, general-audience daily Discord server newspaper from the public activity below. "
        "Do not include sensitive personal information, drama, insults, private details, or anything embarrassing. "
        "Avoid direct quotes unless harmless and public; prefer paraphrase. If activity is low, make a short quiet-day edition. "
        "Use these sections: Headline of the Day, Top Stories, Funniest/Most Notable Moments, Most Active Channels, "
        "Poll or Question of the Day, Upcoming Events or Reminders if available, Quote of the Day only if clearly appropriate, Closing Line.\n\n"
        f"Server: {guild.name}\nTimezone: {timezone_name}\nMost active channels by message count: {active_channels or 'None'}\n\n"
        "Public activity:\n" + "\n".join(activity_lines[-int(module_config.get("max_total_messages", 800)):])
    )

    messages = [
        dict(role="system", content="You are a safe, upbeat, family-friendly community newspaper editor for a Discord server."),
        dict(role="user", content=prompt),
    ]
    return await generate_module_text(loaded_config, module_model(loaded_config, module_config, channel), messages)


async def maybe_post_newspaper_for_guild(guild: discord.Guild, loaded_config: dict[str, Any], force: bool = False) -> bool:
    module_config = get_module_config(loaded_config, "server_newspaper", guild.id)
    if not force and not module_config.get("enabled", False):
        return False

    output_channel_id = get_module_setting(guild.id, "server_newspaper", "output_channel_id")
    if not output_channel_id:
        return False

    channel = guild.get_channel(int(output_channel_id))
    if not isinstance(channel, discord.TextChannel):
        return False

    bot_perms = channel.permissions_for(guild.me)
    if not (bot_perms.view_channel and bot_perms.send_messages and bot_perms.read_message_history):
        return False

    timezone_name = module_config.get("timezone", "America/New_York")
    issue_date = newspaper_issue_date(timezone_name)
    if not force and newspaper_already_posted(guild.id, issue_date):
        return False

    newspaper = await generate_newspaper(guild, channel, loaded_config)
    await send_channel_chunks(channel, newspaper)
    record_newspaper_run(guild.id, issue_date, channel.id)
    return True


async def newspaper_scheduler() -> None:
    await discord_bot.wait_until_ready()

    while not discord_bot.is_closed():
        try:
            loaded_config = await asyncio.to_thread(get_config)

            for guild in discord_bot.guilds:
                module_config = get_module_config(loaded_config, "server_newspaper", guild.id)
                if not module_config.get("enabled", False):
                    continue

                timezone_name = module_config.get("timezone", "America/New_York")
                post_time = str(module_config.get("post_time", "00:00"))
                try:
                    tzinfo = ZoneInfo(timezone_name)
                except ZoneInfoNotFoundError:
                    tzinfo = ZoneInfo("America/New_York")

                now_local = datetime.now(tzinfo)
                try:
                    hour, minute = [int(part) for part in post_time.split(":", 1)]
                except ValueError:
                    hour, minute = 0, 0

                scheduled_today = now_local.replace(hour=hour, minute=minute, second=0, microsecond=0)
                if now_local >= scheduled_today:
                    await maybe_post_newspaper_for_guild(guild, loaded_config)
        except Exception:
            logging.exception("Error in newspaper scheduler")

        await asyncio.sleep(60)


def mediator_high_stakes_text(parts: list[str]) -> str:
    return "\n".join(part for part in parts if part)


async def run_mediator(interaction: discord.Interaction, topic: str, side_a: Optional[str] = None, side_b: Optional[str] = None) -> str:
    loaded_config = await asyncio.to_thread(get_config)
    module_config = get_module_config(loaded_config, "mediator", interaction.guild.id if interaction.guild else None)
    if not module_config.get("enabled", True):
        return "The mediator module is disabled here."

    combined = mediator_high_stakes_text([topic, side_a or "", side_b or ""])
    if REFUSE_MEDIATION_RE.search(combined):
        return (
            "I can't mediate situations involving threats, violence, self-harm, abuse, harassment, or emergencies. "
            "Please involve server moderators or appropriate real-world help."
        )

    disclaimer = ""
    if SERIOUS_MEDIATION_RE.search(combined):
        disclaimer = (
            "Note: this sounds potentially high-stakes. I can help organize thoughts, but this is not legal, medical, "
            "mental health, emergency, or professional advice. Consider involving a qualified person or moderator.\n\n"
        )

    prompt = (
        "Mediate this Discord disagreement neutrally. Do not pick a winner unless explicitly asked. "
        "Use this format: Neutral summary, Side A's likely concern, Side B's likely concern, Common ground, "
        "Suggested compromise, Next step. Keep it calm, non-judgmental, and practical.\n\n"
        f"Topic: {topic}\nSide A: {side_a or 'Not provided'}\nSide B: {side_b or 'Not provided'}"
    )
    messages = [
        dict(role="system", content="You are a calm, neutral mediator for low-stakes Discord disagreements and planning conflicts."),
        dict(role="user", content=prompt),
    ]
    provider_slash_model = module_model(loaded_config, module_config, interaction.channel)
    return disclaimer + await generate_module_text(loaded_config, provider_slash_model, messages)


TITLE_RULES = {
    "Night Owl": ("night_messages", 10),
    "Early Bird": ("early_messages", 10),
    "Link Supplier": ("links", 10),
    "Question Asker": ("questions", 15),
    "Helpful Human": ("thanks_received", 5),
    "Reaction Magnet": ("reactions_received", 25),
    "Conversation Starter": ("conversation_starts", 10),
    "Voice Chat Regular": ("voice_joins", 8),
    "Music Menace": ("music_commands", 5),
    "Poll Goblin": ("poll_activity", 5),
    "Game Goblin": ("party_games", 5),
    "Server Regular": ("active_days", 7),
}

STATIC_PROMPTS = {
    "would_you_rather": [
        ("Have unlimited snacks during movie night", "Have perfect seats at every concert"),
        ("Always win board games", "Always pick the best restaurant"),
        ("Live in a cozy cabin for a month", "Live in a beach house for a month"),
    ],
    "never_have_i_ever": [
        "forgotten why I walked into a room",
        "sent a message to the wrong chat",
        "started a game backlog and made it worse",
        "laughed at my own typo",
    ],
    "this_or_that": [
        ("Pizza night", "Taco night"),
        ("Co-op games", "Competitive games"),
        ("Rainy day playlist", "Sunny day playlist"),
    ],
    "trivia": [
        ("general", "easy", "What planet is known as the Red Planet?", ["Mars", "Venus", "Jupiter", "Mercury"], 0),
        ("gaming", "easy", "Which game features blocks, crafting, and Creepers?", ["Minecraft", "Stardew Valley", "Portal", "Hades"], 0),
        ("science", "easy", "What gas do plants absorb from the air?", ["Carbon dioxide", "Helium", "Oxygen", "Neon"], 0),
        ("movies", "easy", "What is the name of the toy cowboy in Toy Story?", ["Woody", "Buzz", "Andy", "Rex"], 0),
    ],
    "guess": [
        ("object", "backpack"),
        ("place", "library"),
        ("thing", "headphones"),
        ("character", "Mario"),
    ],
}

SAFE_GAME_RE = re.compile(r"\b(sex|sexual|kill|murder|suicide|self[- ]?harm|hate|slur|drug|illegal|abuse|harass|explicit|nsfw)\b", re.IGNORECASE)


def reputation_enabled(guild_id: int, loaded_config: dict[str, Any]) -> bool:
    return bool(get_module_config(loaded_config, "reputation_titles", guild_id).get("enabled", True))


def reputation_is_opted_out(guild_id: int, user_id: int) -> bool:
    with get_db() as conn:
        row = conn.execute(
            "SELECT opted_out_at FROM reputation_title_profiles WHERE guild_id = ? AND user_id = ?",
            (guild_id, user_id),
        ).fetchone()
    return bool(row and row["opted_out_at"])


def set_reputation_optout(guild_id: int, user_id: int, opted_out: bool) -> None:
    with get_db() as conn:
        if opted_out:
            conn.execute(
                "INSERT INTO reputation_title_profiles (guild_id, user_id, opted_out_at) VALUES (?, ?, ?) "
                "ON CONFLICT(guild_id, user_id) DO UPDATE SET opted_out_at = excluded.opted_out_at",
                (guild_id, user_id, now_iso()),
            )
        else:
            conn.execute(
                "INSERT INTO reputation_title_profiles (guild_id, user_id, opted_out_at) VALUES (?, ?, NULL) "
                "ON CONFLICT(guild_id, user_id) DO UPDATE SET opted_out_at = NULL",
                (guild_id, user_id),
            )


def increment_reputation_stat(guild_id: int, user_id: int, stat: str, amount: int = 1) -> int:
    with get_db() as conn:
        conn.execute(
            "INSERT INTO reputation_title_stats (guild_id, user_id, stat, value, updated_at) VALUES (?, ?, ?, ?, ?) "
            "ON CONFLICT(guild_id, user_id, stat) DO UPDATE SET value = value + excluded.value, updated_at = excluded.updated_at",
            (guild_id, user_id, stat, amount, now_iso()),
        )
        row = conn.execute(
            "SELECT value FROM reputation_title_stats WHERE guild_id = ? AND user_id = ? AND stat = ?",
            (guild_id, user_id, stat),
        ).fetchone()
    return int(row["value"] if row else 0)


def grant_reputation_title(guild_id: int, user_id: int, title: str, source: str = "auto", granted_by: Optional[int] = None) -> bool:
    with get_db() as conn:
        existing = conn.execute(
            "SELECT 1 FROM reputation_title_user_titles WHERE guild_id = ? AND user_id = ? AND title = ?",
            (guild_id, user_id, title),
        ).fetchone()
        if existing:
            return False
        conn.execute(
            "INSERT INTO reputation_title_user_titles (guild_id, user_id, title, earned_at, granted_by, source) VALUES (?, ?, ?, ?, ?, ?)",
            (guild_id, user_id, title, now_iso(), granted_by, source),
        )
    return True


def remove_reputation_title(guild_id: int, user_id: int, title: str) -> None:
    with get_db() as conn:
        conn.execute(
            "DELETE FROM reputation_title_user_titles WHERE guild_id = ? AND user_id = ? AND title = ?",
            (guild_id, user_id, title),
        )
        conn.execute(
            "UPDATE reputation_title_profiles SET equipped_title = NULL WHERE guild_id = ? AND user_id = ? AND equipped_title = ?",
            (guild_id, user_id, title),
        )


def maybe_award_reputation_titles(guild_id: int, user_id: int, stat: str, value: int) -> list[str]:
    awarded = []
    for title, (rule_stat, threshold) in TITLE_RULES.items():
        if rule_stat == stat and value >= threshold and grant_reputation_title(guild_id, user_id, title):
            awarded.append(title)
    return awarded


async def track_reputation_stat(guild_id: int, user_id: int, stat: str, loaded_config: dict[str, Any], amount: int = 1, cooldown_key: Optional[tuple[Any, ...]] = None) -> None:
    if not reputation_enabled(guild_id, loaded_config) or reputation_is_opted_out(guild_id, user_id):
        return

    module_config = get_module_config(loaded_config, "reputation_titles", guild_id)
    cooldown_seconds = int(module_config.get("spam_cooldown_seconds", 60))
    if cooldown_key:
        now_ts = datetime.now().timestamp()
        key = (guild_id, user_id, stat, *cooldown_key)
        if reputation_cooldowns.get(key, 0) > now_ts:
            return
        reputation_cooldowns[key] = now_ts + cooldown_seconds

    value = increment_reputation_stat(guild_id, user_id, stat, amount)
    maybe_award_reputation_titles(guild_id, user_id, stat, value)


async def track_reputation_message(message: discord.Message, loaded_config: dict[str, Any]) -> None:
    if message.guild is None or message.author.bot:
        return

    module_config = get_module_config(loaded_config, "reputation_titles", message.guild.id)
    if not module_config.get("track_messages", True):
        return
    if not is_public_server_text_channel(message.channel, set(int(id) for id in module_config.get("ignored_channel_ids", []) or [])):
        return

    await track_reputation_stat(message.guild.id, message.author.id, "messages", loaded_config, cooldown_key=(message.channel.id,))
    today_stat = f"active_day:{datetime.now().astimezone().date().isoformat()}"
    if increment_reputation_stat(message.guild.id, message.author.id, today_stat, 1) == 1:
        value = increment_reputation_stat(message.guild.id, message.author.id, "active_days", 1)
        maybe_award_reputation_titles(message.guild.id, message.author.id, "active_days", value)

    content = message.content or ""
    local_hour = datetime.now().astimezone().hour
    if local_hour >= 23 or local_hour < 5:
        await track_reputation_stat(message.guild.id, message.author.id, "night_messages", loaded_config, cooldown_key=("night",))
    if 5 <= local_hour < 9:
        await track_reputation_stat(message.guild.id, message.author.id, "early_messages", loaded_config, cooldown_key=("early",))
    if URL_RE.search(content):
        await track_reputation_stat(message.guild.id, message.author.id, "links", loaded_config, cooldown_key=("link",))
    if "?" in content:
        await track_reputation_stat(message.guild.id, message.author.id, "questions", loaded_config, cooldown_key=("question",))
    if message.reference is None and len(content) > 20:
        await track_reputation_stat(message.guild.id, message.author.id, "conversation_starts", loaded_config, cooldown_key=("starter", message.channel.id))
    if "poll" in content.lower():
        await track_reputation_stat(message.guild.id, message.author.id, "poll_activity", loaded_config, cooldown_key=("poll",))


def user_titles(guild_id: int, user_id: int) -> tuple[list[str], Optional[str]]:
    with get_db() as conn:
        rows = conn.execute(
            "SELECT title FROM reputation_title_user_titles WHERE guild_id = ? AND user_id = ? ORDER BY title",
            (guild_id, user_id),
        ).fetchall()
        profile = conn.execute(
            "SELECT equipped_title FROM reputation_title_profiles WHERE guild_id = ? AND user_id = ?",
            (guild_id, user_id),
        ).fetchone()
    return [row["title"] for row in rows], (profile["equipped_title"] if profile else None)


def party_enabled(guild_id: int, loaded_config: dict[str, Any], game: Optional[str] = None) -> bool:
    module_config = get_module_config(loaded_config, "party_games", guild_id)
    if not module_config.get("enabled", True):
        return False
    if game:
        return bool((module_config.get("enabled_games") or {}).get(game, True))
    return True


def party_on_cooldown(guild_id: int, user_id: int, loaded_config: dict[str, Any]) -> Optional[int]:
    module_config = get_module_config(loaded_config, "party_games", guild_id)
    cooldown_seconds = int(module_config.get("cooldown_seconds", 20))
    key = (guild_id, user_id)
    now_ts = datetime.now().timestamp()
    if party_cooldowns.get(key, 0) > now_ts:
        return int(party_cooldowns[key] - now_ts)
    party_cooldowns[key] = now_ts + cooldown_seconds
    return None


def add_party_score(guild_id: int, user_id: int, game: str, points: int = 1, streak_delta: int = 1) -> None:
    with get_db() as conn:
        conn.execute(
            "INSERT INTO party_game_scores (guild_id, user_id, game, score, streak, updated_at) VALUES (?, ?, ?, ?, ?, ?) "
            "ON CONFLICT(guild_id, user_id, game) DO UPDATE SET score = score + excluded.score, streak = MAX(0, streak + excluded.streak), updated_at = excluded.updated_at",
            (guild_id, user_id, game, points, streak_delta, now_iso()),
        )
        conn.execute(
            "INSERT INTO reputation_title_stats (guild_id, user_id, stat, value, updated_at) VALUES (?, ?, 'party_games', 1, ?) "
            "ON CONFLICT(guild_id, user_id, stat) DO UPDATE SET value = value + 1, updated_at = excluded.updated_at",
            (guild_id, user_id, now_iso()),
        )
        row = conn.execute(
            "SELECT value FROM reputation_title_stats WHERE guild_id = ? AND user_id = ? AND stat = 'party_games'",
            (guild_id, user_id),
        ).fetchone()
    if row:
        maybe_award_reputation_titles(guild_id, user_id, "party_games", int(row["value"]))


def record_party_history(guild_id: int, channel_id: int, game: str, user_id: Optional[int], result: str) -> None:
    with get_db() as conn:
        conn.execute(
            "INSERT INTO party_game_history (guild_id, channel_id, game, user_id, result, created_at) VALUES (?, ?, ?, ?, ?, ?)",
            (guild_id, channel_id, game, user_id, result, now_iso()),
        )


async def safe_ai_game_prompt(interaction: discord.Interaction, game: str, loaded_config: dict[str, Any], fallback: Any) -> Any:
    module_config = get_module_config(loaded_config, "party_games", interaction.guild.id)
    if not module_config.get("ai_generated_prompts", True):
        return fallback
    prompt = (
        f"Generate one safe, general-audience Discord party game prompt for {game}. "
        "Avoid sexual, hateful, violent, self-harm, illegal, harassment, invasive, or embarrassing content. "
        "Return compact JSON only."
    )
    try:
        text = await generate_module_text(
            loaded_config,
            module_model(loaded_config, module_config, interaction.channel),
            [dict(role="system", content="You generate safe party game prompts as JSON only."), dict(role="user", content=prompt)],
        )
        if SAFE_GAME_RE.search(text):
            return fallback
        parsed = json_loads(re.search(r"\{[\s\S]*\}|\[[\s\S]*\]", text).group(0) if re.search(r"\{[\s\S]*\}|\[[\s\S]*\]", text) else text, None)
        return parsed or fallback
    except Exception:
        logging.exception("AI party prompt generation failed")
        return fallback


class VoteView(discord.ui.View):
    def __init__(self, labels: list[str], timeout: int = 60):
        super().__init__(timeout=timeout)
        self.votes: dict[int, int] = {}
        for index, label in enumerate(labels):
            self.add_item(VoteButton(label[:80], index))

    def results(self) -> list[int]:
        counts = [0 for _ in self.children]
        for vote in self.votes.values():
            if 0 <= vote < len(counts):
                counts[vote] += 1
        return counts


class VoteButton(discord.ui.Button):
    def __init__(self, label: str, index: int):
        super().__init__(label=label, style=discord.ButtonStyle.primary)
        self.index = index

    async def callback(self, interaction: discord.Interaction) -> None:
        view = self.view
        if isinstance(view, VoteView):
            view.votes[interaction.user.id] = self.index
        await interaction.response.send_message("Vote counted.", ephemeral=True)


class TriviaView(discord.ui.View):
    def __init__(self, guild_id: int, game: str, answers: list[str], correct_index: int, timeout: int = 60):
        super().__init__(timeout=timeout)
        self.guild_id = guild_id
        self.game = game
        self.correct_index = correct_index
        self.answered: set[int] = set()
        for index, answer in enumerate(answers):
            self.add_item(TriviaButton(answer[:80], index))


class TriviaButton(discord.ui.Button):
    def __init__(self, label: str, index: int):
        super().__init__(label=label, style=discord.ButtonStyle.secondary)
        self.index = index

    async def callback(self, interaction: discord.Interaction) -> None:
        view = self.view
        if not isinstance(view, TriviaView):
            return
        if interaction.user.id in view.answered:
            await interaction.response.send_message("You already answered.", ephemeral=True)
            return
        view.answered.add(interaction.user.id)
        if self.index == view.correct_index:
            add_party_score(view.guild_id, interaction.user.id, view.game, 1, 1)
            await interaction.response.send_message("Correct!", ephemeral=True)
        else:
            add_party_score(view.guild_id, interaction.user.id, view.game, 0, -1)
            await interaction.response.send_message("Not quite.", ephemeral=True)


def attachment_extension(attachment: discord.Attachment) -> str:
    return attachment.filename.rsplit(".", 1)[-1].lower() if "." in attachment.filename else ""


def attachment_is_image(attachment: discord.Attachment) -> bool:
    return (attachment.content_type or "").lower().startswith("image/") or attachment_extension(attachment) in {"png", "jpg", "jpeg", "webp"}


async def extract_attachment_for_brain(attachment: discord.Attachment, module_config: dict[str, Any]) -> tuple[str, list[dict[str, Any]], list[str]]:
    warnings = []
    max_file_size = int(float(module_config.get("max_file_size_mb", 10)) * 1024 * 1024)
    max_chars = int(module_config.get("max_extracted_chars", 50000))
    extension = attachment_extension(attachment)
    allowed_extensions = set(module_config.get("allowed_extensions", []) or [])

    if extension not in allowed_extensions:
        raise ValueError(f"Unsupported file type: .{extension or 'unknown'}")
    if attachment.size and attachment.size > max_file_size:
        raise ValueError(f"File is too large. Limit is {module_config.get('max_file_size_mb', 10)} MB.")

    response = await httpx_client.get(attachment.url)
    response.raise_for_status()

    if attachment_is_image(attachment):
        return "", [dict(type="image_url", image_url=dict(url=f"data:{attachment.content_type};base64,{b64encode(response.content).decode('utf-8')}"))], warnings

    kind = get_attachment_kind(attachment)
    if kind == "pdf":
        text = await asyncio.to_thread(extract_pdf_text, response.content)
    elif kind == "docx":
        text = await asyncio.to_thread(extract_docx_text, response.content)
    else:
        text = response.text

    text = normalize_extracted_text(text)
    if len(text) > max_chars:
        text = text[:max_chars]
        warnings.append(f"Extracted text was truncated to {max_chars:,} characters.")
    return text, [], warnings


async def run_attachment_brain(interaction: discord.Interaction, attachment: discord.Attachment, task: str, question: Optional[str] = None) -> str:
    loaded_config = await asyncio.to_thread(get_config)
    module_config = get_module_config(loaded_config, "attachment_brain", interaction.guild.id if interaction.guild else None)
    if not module_config.get("enabled", True):
        return "Attachment Brain is disabled."

    text, images, warnings = await extract_attachment_for_brain(attachment, module_config)
    provider_slash_model = module_config.get("model") if module_config.get("model") in loaded_config["models"] else get_effective_model(interaction.channel, loaded_config)

    if images:
        vision_model = module_config.get("vision_model")
        if vision_model in loaded_config["models"]:
            provider_slash_model = vision_model
        elif not any(tag in provider_slash_model.lower() for tag in VISION_MODEL_TAGS):
            return "Image understanding requires a vision-capable model. Set `modules.attachment_brain.vision_model` or use a vision model in this channel."

    if task == "extract-text":
        if not text:
            return "No readable text was extracted."
        preview = text if len(text) <= 1800 else text[:1800] + "\n\n...[truncated preview]"
        return "\n".join(warnings + [preview])

    instruction = {
        "summarize": "Summarize the main points of this file.",
        "ask": f"Answer this question based only on the file. If the answer is not in the file, say so.\nQuestion: {question}",
        "debug-log": "Analyze this log. Identify likely causes, important errors, and practical fixes.",
        "explain-code": "Explain what this code does, identify obvious bugs or risky parts, and suggest improvements. Do not execute it.",
        "convert-json": "Convert or normalize the file content into clear JSON if reasonable. If not possible, explain why.",
    }.get(task, "Analyze this file.")

    content = [dict(type="text", text=f"File: {attachment.filename}\nTask: {instruction}\n\nExtracted text:\n{text or '[image attached]'}")] + images
    messages = [
        dict(role="system", content="You answer questions about uploaded files safely. Do not execute untrusted code."),
        dict(role="user", content=content if images else content[0]["text"]),
    ]
    answer = await generate_module_text(loaded_config, provider_slash_model, messages)
    return "\n".join(warnings + [answer])


newspaper_group = discord.app_commands.Group(name="newspaper", description="Daily server newspaper controls")
mediator_group = discord.app_commands.Group(name="mediator", description="AI mediator module controls")
titles_group = discord.app_commands.Group(name="titles", description="Reputation title profiles")
party_group = discord.app_commands.Group(name="party", description="General-audience party games")
file_group = discord.app_commands.Group(name="file", description="Ask questions about uploaded files")


@newspaper_group.command(name="generate", description="Generate today's server newspaper now")
async def newspaper_generate(interaction: discord.Interaction) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not await require_admin_interaction(interaction, loaded_config):
        return

    await interaction.response.defer(thinking=True)
    posted = await maybe_post_newspaper_for_guild(interaction.guild, loaded_config, force=True)
    await interaction.followup.send("Newspaper generated." if posted else "I could not generate the newspaper. Check the output channel and permissions.", ephemeral=True)


@newspaper_group.command(name="set-channel", description="Set the newspaper output channel")
async def newspaper_set_channel(interaction: discord.Interaction, channel: discord.TextChannel) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not await require_admin_interaction(interaction, loaded_config):
        return

    set_module_setting(interaction.guild.id, "server_newspaper", "output_channel_id", channel.id)
    await interaction.response.send_message(f"Newspaper output channel set to {channel.mention}.", ephemeral=True)


@newspaper_group.command(name="enable", description="Enable the daily server newspaper")
async def newspaper_enable(interaction: discord.Interaction) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not await require_admin_interaction(interaction, loaded_config):
        return

    set_module_setting(interaction.guild.id, "server_newspaper", "enabled", True)
    await interaction.response.send_message("Daily newspaper enabled.", ephemeral=True)


@newspaper_group.command(name="disable", description="Disable the daily server newspaper")
async def newspaper_disable(interaction: discord.Interaction) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not await require_admin_interaction(interaction, loaded_config):
        return

    set_module_setting(interaction.guild.id, "server_newspaper", "enabled", False)
    await interaction.response.send_message("Daily newspaper disabled.", ephemeral=True)


@newspaper_group.command(name="ignore-channel", description="Exclude a channel from newspaper summaries")
async def newspaper_ignore_channel(interaction: discord.Interaction, channel: discord.TextChannel) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not await require_admin_interaction(interaction, loaded_config):
        return

    set_ignored_channel(interaction.guild.id, "server_newspaper", channel.id, True)
    await interaction.response.send_message(f"{channel.mention} will be ignored by the newspaper.", ephemeral=True)


@newspaper_group.command(name="unignore-channel", description="Include a channel in newspaper summaries again")
async def newspaper_unignore_channel(interaction: discord.Interaction, channel: discord.TextChannel) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not await require_admin_interaction(interaction, loaded_config):
        return

    set_ignored_channel(interaction.guild.id, "server_newspaper", channel.id, False)
    await interaction.response.send_message(f"{channel.mention} can be included in the newspaper again.", ephemeral=True)


@newspaper_group.command(name="status", description="Show newspaper settings")
async def newspaper_status(interaction: discord.Interaction) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if interaction.guild is None:
        await interaction.response.send_message("This command only works in a server.", ephemeral=True)
        return

    module_config = get_module_config(loaded_config, "server_newspaper", interaction.guild.id)
    output_channel_id = get_module_setting(interaction.guild.id, "server_newspaper", "output_channel_id")
    output_channel = interaction.guild.get_channel(int(output_channel_id)) if output_channel_id else None
    issue_date = newspaper_issue_date(module_config.get("timezone", "America/New_York"))

    await interaction.response.send_message(
        "\n".join(
            [
                f"Enabled: `{bool(module_config.get('enabled', False))}`",
                f"Output channel: {output_channel.mention if output_channel else '`not set`'}",
                f"Timezone: `{module_config.get('timezone', 'America/New_York')}`",
                f"Post time: `{module_config.get('post_time', '00:00')}`",
                f"Posted today: `{newspaper_already_posted(interaction.guild.id, issue_date)}`",
                f"Ignored channels: `{len(module_config.get('ignored_channel_ids', []))}`",
            ]
        ),
        ephemeral=True,
    )


@discord.app_commands.allowed_installs(guilds=True, users=False)
@discord.app_commands.allowed_contexts(guilds=True, dms=False, private_channels=False)
@discord.app_commands.describe(topic="The disagreement, planning conflict, or decision to mediate", side_a="Optional side A", side_b="Optional side B")
@discord_bot.tree.command(name="mediate", description="Get a neutral AI mediation summary")
async def mediate_command(interaction: discord.Interaction, topic: str, side_a: Optional[str] = None, side_b: Optional[str] = None) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not user_has_permission_for_interaction(interaction, loaded_config):
        await interaction.response.send_message("You don't have permission to use this bot here.", ephemeral=True)
        return

    module_config = get_module_config(loaded_config, "mediator", interaction.guild.id if interaction.guild else None)
    cooldown_seconds = int(module_config.get("cooldown_seconds", 60))
    cooldown_key = (interaction.guild.id if interaction.guild else 0, interaction.user.id)
    now_ts = datetime.now().timestamp()
    next_allowed = mediator_cooldowns.get(cooldown_key, 0)
    if now_ts < next_allowed:
        await interaction.response.send_message(f"Slow down a little. Try again in {int(next_allowed - now_ts)} seconds.", ephemeral=True)
        return

    mediator_cooldowns[cooldown_key] = now_ts + cooldown_seconds
    await interaction.response.defer(thinking=True)
    output = await run_mediator(interaction, topic, side_a, side_b)
    await send_interaction_chunks(interaction, output, private=False)


@mediator_group.command(name="enable", description="Enable the mediator module")
async def mediator_enable(interaction: discord.Interaction) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not await require_admin_interaction(interaction, loaded_config):
        return
    set_module_setting(interaction.guild.id, "mediator", "enabled", True)
    await interaction.response.send_message("Mediator enabled.", ephemeral=True)


@mediator_group.command(name="disable", description="Disable the mediator module")
async def mediator_disable(interaction: discord.Interaction) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not await require_admin_interaction(interaction, loaded_config):
        return
    set_module_setting(interaction.guild.id, "mediator", "enabled", False)
    await interaction.response.send_message("Mediator disabled.", ephemeral=True)


@mediator_group.command(name="status", description="Show mediator module settings")
async def mediator_status(interaction: discord.Interaction) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if interaction.guild is None:
        await interaction.response.send_message("This command only works in a server.", ephemeral=True)
        return
    module_config = get_module_config(loaded_config, "mediator", interaction.guild.id)
    await interaction.response.send_message(
        f"Enabled: `{bool(module_config.get('enabled', True))}`\nCooldown: `{module_config.get('cooldown_seconds', 60)}s`\nModel: `{module_config.get('model') or 'channel default'}`",
        ephemeral=True,
    )


@titles_group.command(name="profile", description="View your reputation title profile")
async def titles_profile(interaction: discord.Interaction, user: Optional[discord.Member] = None) -> None:
    if interaction.guild is None:
        await interaction.response.send_message("This command only works in a server.", ephemeral=True)
        return
    target = user or interaction.user
    titles, equipped = user_titles(interaction.guild.id, target.id)
    with get_db() as conn:
        stats = conn.execute(
            "SELECT stat, value FROM reputation_title_stats WHERE guild_id = ? AND user_id = ? ORDER BY value DESC LIMIT 8",
            (interaction.guild.id, target.id),
        ).fetchall()
    stat_text = ", ".join(f"{row['stat']}: {row['value']}" for row in stats) or "No tracked stats yet."
    title_text = ", ".join(titles) or "No titles yet."
    await interaction.response.send_message(
        f"Profile for {target.mention}\nEquipped: `{equipped or 'none'}`\nTitles: {title_text}\nStats: {stat_text}",
        ephemeral=user is None,
    )


@titles_group.command(name="list", description="List available reputation titles")
async def titles_list(interaction: discord.Interaction) -> None:
    lines = [f"- **{title}**: `{stat}` >= {threshold}" for title, (stat, threshold) in TITLE_RULES.items()]
    await interaction.response.send_message("\n".join(lines), ephemeral=True)


@titles_group.command(name="equip", description="Equip one of your earned titles")
async def titles_equip(interaction: discord.Interaction, title: str) -> None:
    if interaction.guild is None:
        await interaction.response.send_message("This command only works in a server.", ephemeral=True)
        return
    titles, _ = user_titles(interaction.guild.id, interaction.user.id)
    if title not in titles:
        await interaction.response.send_message("You have not earned that title.", ephemeral=True)
        return
    with get_db() as conn:
        conn.execute(
            "INSERT INTO reputation_title_profiles (guild_id, user_id, equipped_title) VALUES (?, ?, ?) "
            "ON CONFLICT(guild_id, user_id) DO UPDATE SET equipped_title = excluded.equipped_title",
            (interaction.guild.id, interaction.user.id, title),
        )
    await interaction.response.send_message(f"Equipped title: **{title}**", ephemeral=True)


@titles_group.command(name="leaderboard", description="Show who has the most titles")
async def titles_leaderboard(interaction: discord.Interaction) -> None:
    if interaction.guild is None:
        await interaction.response.send_message("This command only works in a server.", ephemeral=True)
        return
    with get_db() as conn:
        rows = conn.execute(
            "SELECT user_id, COUNT(*) AS count FROM reputation_title_user_titles WHERE guild_id = ? GROUP BY user_id ORDER BY count DESC LIMIT 10",
            (interaction.guild.id,),
        ).fetchall()
    if not rows:
        await interaction.response.send_message("No titles have been earned yet.")
        return
    await interaction.response.send_message("\n".join(f"{i}. <@{row['user_id']}> - {row['count']} titles" for i, row in enumerate(rows, 1)))


@titles_group.command(name="opt-out", description="Opt out of reputation title tracking")
async def titles_opt_out(interaction: discord.Interaction) -> None:
    if interaction.guild is None:
        await interaction.response.send_message("This command only works in a server.", ephemeral=True)
        return
    set_reputation_optout(interaction.guild.id, interaction.user.id, True)
    await interaction.response.send_message("You are opted out of reputation title tracking.", ephemeral=True)


@titles_group.command(name="opt-in", description="Opt back into reputation title tracking")
async def titles_opt_in(interaction: discord.Interaction) -> None:
    if interaction.guild is None:
        await interaction.response.send_message("This command only works in a server.", ephemeral=True)
        return
    set_reputation_optout(interaction.guild.id, interaction.user.id, False)
    await interaction.response.send_message("You are opted back into reputation title tracking.", ephemeral=True)


@titles_group.command(name="grant", description="Admin grant a title")
async def titles_grant(interaction: discord.Interaction, user: discord.Member, title: str) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not await require_admin_interaction(interaction, loaded_config):
        return
    granted = grant_reputation_title(interaction.guild.id, user.id, title, source="admin", granted_by=interaction.user.id)
    await interaction.response.send_message(f"{'Granted' if granted else 'Already had'} **{title}** for {user.mention}.", ephemeral=True)


@titles_group.command(name="remove", description="Admin remove a title")
async def titles_remove(interaction: discord.Interaction, user: discord.Member, title: str) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not await require_admin_interaction(interaction, loaded_config):
        return
    remove_reputation_title(interaction.guild.id, user.id, title)
    await interaction.response.send_message(f"Removed **{title}** from {user.mention}.", ephemeral=True)


@titles_group.command(name="status", description="Show reputation title settings")
async def titles_status(interaction: discord.Interaction) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not await require_admin_interaction(interaction, loaded_config):
        return
    module_config = get_module_config(loaded_config, "reputation_titles", interaction.guild.id)
    await interaction.response.send_message(
        f"Enabled: `{module_config.get('enabled', True)}`\n"
        f"Track messages: `{module_config.get('track_messages', True)}`\n"
        f"Track voice: `{module_config.get('track_voice', True)}`\n"
        f"Track reactions: `{module_config.get('track_reactions', True)}`\n"
        f"Spam cooldown: `{module_config.get('spam_cooldown_seconds', 60)}s`",
        ephemeral=True,
    )


@titles_group.command(name="enable", description="Enable reputation title tracking")
async def titles_enable(interaction: discord.Interaction) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not await require_admin_interaction(interaction, loaded_config):
        return
    set_module_setting(interaction.guild.id, "reputation_titles", "enabled", True)
    await interaction.response.send_message("Reputation title tracking enabled.", ephemeral=True)


@titles_group.command(name="disable", description="Disable reputation title tracking")
async def titles_disable(interaction: discord.Interaction) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not await require_admin_interaction(interaction, loaded_config):
        return
    set_module_setting(interaction.guild.id, "reputation_titles", "enabled", False)
    await interaction.response.send_message("Reputation title tracking disabled.", ephemeral=True)


async def run_party_vote(interaction: discord.Interaction, game: str, title: str, labels: list[str]) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if interaction.guild is None or not party_enabled(interaction.guild.id, loaded_config, game):
        await interaction.response.send_message("That party game is disabled here.", ephemeral=True)
        return
    if cooldown := party_on_cooldown(interaction.guild.id, interaction.user.id, loaded_config):
        await interaction.response.send_message(f"Try again in {cooldown} seconds.", ephemeral=True)
        return
    labels = [normalize_extracted_text(str(label))[:80] for label in labels if normalize_extracted_text(str(label)) and not SAFE_GAME_RE.search(str(label))]
    if len(labels) < 2:
        labels = ["Option A", "Option B"]
    timeout = int(get_module_config(loaded_config, "party_games", interaction.guild.id).get("default_round_timeout_seconds", 60))
    view = VoteView(labels, timeout=timeout)
    embed = discord.Embed(title=title, description="\n".join(f"{i + 1}. {label}" for i, label in enumerate(labels)), color=discord.Color.blurple())
    await interaction.response.send_message(embed=embed, view=view)
    await view.wait()
    results = view.results()
    record_party_history(interaction.guild.id, interaction.channel.id, game, interaction.user.id, json_dumps(results))
    add_party_score(interaction.guild.id, interaction.user.id, game, 1, 1)
    await interaction.followup.send("Results: " + " | ".join(f"{labels[i]}: {count}" for i, count in enumerate(results)))


@party_group.command(name="would-you-rather", description="Start a Would You Rather vote")
async def party_would_you_rather(interaction: discord.Interaction) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    fallback = random.choice(STATIC_PROMPTS["would_you_rather"])
    prompt = await safe_ai_game_prompt(interaction, "would_you_rather", loaded_config, {"options": list(fallback)})
    labels = prompt.get("options", list(fallback)) if isinstance(prompt, dict) else list(fallback)
    await run_party_vote(interaction, "would_you_rather", "Would You Rather?", labels[:2])


@party_group.command(name="never-have-i-ever", description="Start a light Never Have I Ever prompt")
async def party_never_have_i_ever(interaction: discord.Interaction) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    fallback = random.choice(STATIC_PROMPTS["never_have_i_ever"])
    prompt = await safe_ai_game_prompt(interaction, "never_have_i_ever", loaded_config, {"prompt": fallback})
    statement = prompt.get("prompt", fallback) if isinstance(prompt, dict) else fallback
    await run_party_vote(interaction, "never_have_i_ever", f"Never have I ever {statement}", ["I have", "I have not"])


@party_group.command(name="this-or-that", description="Start a quick This or That vote")
async def party_this_or_that(interaction: discord.Interaction) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    fallback = random.choice(STATIC_PROMPTS["this_or_that"])
    prompt = await safe_ai_game_prompt(interaction, "this_or_that", loaded_config, {"options": list(fallback)})
    labels = prompt.get("options", list(fallback)) if isinstance(prompt, dict) else list(fallback)
    await run_party_vote(interaction, "this_or_that", "This or That?", labels[:2])


@party_group.command(name="trivia", description="Start a trivia question")
async def party_trivia(interaction: discord.Interaction, category: str = "random", difficulty: str = "easy") -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if interaction.guild is None or not party_enabled(interaction.guild.id, loaded_config, "trivia"):
        await interaction.response.send_message("Trivia is disabled here.", ephemeral=True)
        return
    if cooldown := party_on_cooldown(interaction.guild.id, interaction.user.id, loaded_config):
        await interaction.response.send_message(f"Try again in {cooldown} seconds.", ephemeral=True)
        return
    trivia = random.choice([item for item in STATIC_PROMPTS["trivia"] if category == "random" or item[0] == category] or STATIC_PROMPTS["trivia"])
    _, _, question, answers, correct_index = trivia
    timeout = int(get_module_config(loaded_config, "party_games", interaction.guild.id).get("default_round_timeout_seconds", 60))
    view = TriviaView(interaction.guild.id, "trivia", answers, correct_index, timeout=timeout)
    await interaction.response.send_message(embed=discord.Embed(title="Trivia", description=question), view=view)
    record_party_history(interaction.guild.id, interaction.channel.id, "trivia", interaction.user.id, question)


@party_group.command(name="two-truths-start", description="Start Two Truths and a Lie")
async def party_two_truths_start(interaction: discord.Interaction, statement_1: str, statement_2: str, statement_3: str, lie_number: int) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if interaction.guild is None or not party_enabled(interaction.guild.id, loaded_config, "two_truths_and_lie"):
        await interaction.response.send_message("Two Truths and a Lie is disabled here.", ephemeral=True)
        return
    if lie_number not in (1, 2, 3):
        await interaction.response.send_message("Lie number must be 1, 2, or 3.", ephemeral=True)
        return
    statements = [statement_1, statement_2, statement_3]
    if any(SAFE_GAME_RE.search(statement) for statement in statements):
        await interaction.response.send_message("Keep statements general-audience friendly.", ephemeral=True)
        return
    party_rounds[(interaction.guild.id, interaction.channel.id, "two_truths")] = {"lie": lie_number, "host": interaction.user.id}
    await interaction.response.send_message("Two Truths and a Lie:\n" + "\n".join(f"{i + 1}. {statement}" for i, statement in enumerate(statements)) + "\nUse `/party two-truths-guess lie_number:<1-3>`.")


@party_group.command(name="two-truths-guess", description="Guess the lie")
async def party_two_truths_guess(interaction: discord.Interaction, lie_number: int) -> None:
    if interaction.guild is None:
        await interaction.response.send_message("This command only works in a server.", ephemeral=True)
        return
    state = party_rounds.get((interaction.guild.id, interaction.channel.id, "two_truths"))
    if not state:
        await interaction.response.send_message("No active Two Truths round here.", ephemeral=True)
        return
    if lie_number == state["lie"]:
        add_party_score(interaction.guild.id, interaction.user.id, "two_truths_and_lie", 1, 1)
        await interaction.response.send_message("Correct!")
    else:
        await interaction.response.send_message("Nope.", ephemeral=True)


@party_group.command(name="wordchain-start", description="Start a Word Chain game")
async def party_wordchain_start(interaction: discord.Interaction, first_word: str) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if interaction.guild is None or not party_enabled(interaction.guild.id, loaded_config, "word_chain"):
        await interaction.response.send_message("Word Chain is disabled here.", ephemeral=True)
        return
    word = re.sub(r"[^a-zA-Z]", "", first_word).lower()
    if not word:
        await interaction.response.send_message("Start with a word.", ephemeral=True)
        return
    party_rounds[(interaction.guild.id, interaction.channel.id, "word_chain")] = {"last": word, "used": {word}}
    await interaction.response.send_message(f"Word Chain started with **{word}**. Next word must start with `{word[-1]}`. Use `/party wordchain word:<word>`.")


@party_group.command(name="wordchain", description="Play a word in Word Chain")
async def party_wordchain(interaction: discord.Interaction, word: str) -> None:
    if interaction.guild is None:
        await interaction.response.send_message("This command only works in a server.", ephemeral=True)
        return
    state = party_rounds.get((interaction.guild.id, interaction.channel.id, "word_chain"))
    if not state:
        await interaction.response.send_message("No active Word Chain game here.", ephemeral=True)
        return
    cleaned = re.sub(r"[^a-zA-Z]", "", word).lower()
    if not cleaned or cleaned in state["used"] or cleaned[0] != state["last"][-1]:
        await interaction.response.send_message(f"Invalid word. It must start with `{state['last'][-1]}` and not be repeated.", ephemeral=True)
        return
    state["last"] = cleaned
    state["used"].add(cleaned)
    add_party_score(interaction.guild.id, interaction.user.id, "word_chain", 1, 1)
    await interaction.response.send_message(f"Accepted: **{cleaned}**. Next starts with `{cleaned[-1]}`.")


@party_group.command(name="guess-start", description="Start a yes/no guessing game")
async def party_guess_start(interaction: discord.Interaction) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if interaction.guild is None or not party_enabled(interaction.guild.id, loaded_config, "guess"):
        await interaction.response.send_message("Guess is disabled here.", ephemeral=True)
        return
    category, answer = random.choice(STATIC_PROMPTS["guess"])
    party_rounds[(interaction.guild.id, interaction.channel.id, "guess")] = {"answer": answer.lower(), "display": answer, "questions": 0}
    await interaction.response.send_message(f"I am thinking of a safe **{category}**. Use `/party guess-answer guess:<text>`.")


@party_group.command(name="guess-answer", description="Make a guess in the Guess game")
async def party_guess_answer(interaction: discord.Interaction, guess: str) -> None:
    if interaction.guild is None:
        await interaction.response.send_message("This command only works in a server.", ephemeral=True)
        return
    state = party_rounds.get((interaction.guild.id, interaction.channel.id, "guess"))
    if not state:
        await interaction.response.send_message("No active Guess game here.", ephemeral=True)
        return
    if guess.lower().strip() == state["answer"]:
        add_party_score(interaction.guild.id, interaction.user.id, "guess", 1, 1)
        party_rounds.pop((interaction.guild.id, interaction.channel.id, "guess"), None)
        await interaction.response.send_message(f"Correct! It was **{state['display']}**.")
    else:
        await interaction.response.send_message("Not it.", ephemeral=True)


@party_group.command(name="leaderboard", description="Show party game leaderboard")
async def party_leaderboard(interaction: discord.Interaction) -> None:
    if interaction.guild is None:
        await interaction.response.send_message("This command only works in a server.", ephemeral=True)
        return
    with get_db() as conn:
        rows = conn.execute(
            "SELECT user_id, SUM(score) AS total FROM party_game_scores WHERE guild_id = ? GROUP BY user_id ORDER BY total DESC LIMIT 10",
            (interaction.guild.id,),
        ).fetchall()
    await interaction.response.send_message("\n".join(f"{i}. <@{row['user_id']}> - {row['total']}" for i, row in enumerate(rows, 1)) or "No scores yet.")


@party_group.command(name="stats", description="Show your party game stats")
async def party_stats(interaction: discord.Interaction, user: Optional[discord.Member] = None) -> None:
    if interaction.guild is None:
        await interaction.response.send_message("This command only works in a server.", ephemeral=True)
        return
    target = user or interaction.user
    with get_db() as conn:
        rows = conn.execute(
            "SELECT game, score, streak FROM party_game_scores WHERE guild_id = ? AND user_id = ? ORDER BY score DESC",
            (interaction.guild.id, target.id),
        ).fetchall()
    await interaction.response.send_message("\n".join(f"{row['game']}: {row['score']} points, streak {row['streak']}" for row in rows) or "No stats yet.", ephemeral=user is None)


@party_group.command(name="stop", description="Stop party games in this channel")
async def party_stop(interaction: discord.Interaction) -> None:
    if interaction.guild is None:
        await interaction.response.send_message("This command only works in a server.", ephemeral=True)
        return
    for key in list(party_rounds):
        if key[0] == interaction.guild.id and key[1] == interaction.channel.id:
            party_rounds.pop(key, None)
    await interaction.response.send_message("Stopped active party games in this channel.")


@party_group.command(name="status", description="Show party game module status")
async def party_status(interaction: discord.Interaction) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not await require_admin_interaction(interaction, loaded_config):
        return
    module_config = get_module_config(loaded_config, "party_games", interaction.guild.id)
    await interaction.response.send_message(f"Enabled: `{module_config.get('enabled', True)}`\nGames: `{module_config.get('enabled_games')}`", ephemeral=True)


@party_group.command(name="enable", description="Enable a party game")
async def party_enable(interaction: discord.Interaction, game: str) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not await require_admin_interaction(interaction, loaded_config):
        return
    enabled_games = dict(get_module_config(loaded_config, "party_games", interaction.guild.id).get("enabled_games") or {})
    enabled_games[game] = True
    set_module_setting(interaction.guild.id, "party_games", "enabled_games", enabled_games)
    await interaction.response.send_message(f"Enabled `{game}`.", ephemeral=True)


@party_group.command(name="disable", description="Disable a party game")
async def party_disable(interaction: discord.Interaction, game: str) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not await require_admin_interaction(interaction, loaded_config):
        return
    enabled_games = dict(get_module_config(loaded_config, "party_games", interaction.guild.id).get("enabled_games") or {})
    enabled_games[game] = False
    set_module_setting(interaction.guild.id, "party_games", "enabled_games", enabled_games)
    await interaction.response.send_message(f"Disabled `{game}`.", ephemeral=True)


async def file_command_response(interaction: discord.Interaction, attachment: discord.Attachment, task: str, question: Optional[str] = None) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not user_has_permission_for_interaction(interaction, loaded_config):
        await interaction.response.send_message("You don't have permission to use this bot here.", ephemeral=True)
        return
    await interaction.response.defer(thinking=True, ephemeral=False)
    try:
        output = await run_attachment_brain(interaction, attachment, task, question)
        await send_interaction_chunks(interaction, output, private=False)
    except Exception as exc:
        logging.exception("Attachment Brain failed")
        await interaction.followup.send(f"Could not process that file: {exc}", ephemeral=True)


@file_group.command(name="summarize", description="Summarize an uploaded file")
async def file_summarize(interaction: discord.Interaction, attachment: discord.Attachment) -> None:
    await file_command_response(interaction, attachment, "summarize")


@file_group.command(name="ask", description="Ask a question about an uploaded file")
async def file_ask(interaction: discord.Interaction, attachment: discord.Attachment, question: str) -> None:
    await file_command_response(interaction, attachment, "ask", question)


@file_group.command(name="extract-text", description="Extract readable text from a file")
async def file_extract_text(interaction: discord.Interaction, attachment: discord.Attachment) -> None:
    await file_command_response(interaction, attachment, "extract-text")


@file_group.command(name="debug-log", description="Analyze a log file")
async def file_debug_log(interaction: discord.Interaction, attachment: discord.Attachment) -> None:
    await file_command_response(interaction, attachment, "debug-log")


@file_group.command(name="explain-code", description="Explain a code file")
async def file_explain_code(interaction: discord.Interaction, attachment: discord.Attachment) -> None:
    await file_command_response(interaction, attachment, "explain-code")


@file_group.command(name="convert-json", description="Convert or normalize file content to JSON")
async def file_convert_json(interaction: discord.Interaction, attachment: discord.Attachment) -> None:
    await file_command_response(interaction, attachment, "convert-json")


@file_group.command(name="status", description="Show Attachment Brain settings")
async def file_status(interaction: discord.Interaction) -> None:
    loaded_config = await asyncio.to_thread(get_config)
    if not await require_admin_interaction(interaction, loaded_config):
        return
    module_config = get_module_config(loaded_config, "attachment_brain", interaction.guild.id if interaction.guild else None)
    await interaction.response.send_message(
        f"Enabled: `{module_config.get('enabled', True)}`\nMax file size: `{module_config.get('max_file_size_mb', 10)} MB`\nMax extracted chars: `{module_config.get('max_extracted_chars', 50000)}`",
        ephemeral=True,
    )


discord_bot.tree.add_command(newspaper_group)
discord_bot.tree.add_command(mediator_group)
discord_bot.tree.add_command(titles_group)
discord_bot.tree.add_command(party_group)
discord_bot.tree.add_command(file_group)


@discord.app_commands.allowed_installs(guilds=True, users=True)
@discord.app_commands.allowed_contexts(guilds=True, dms=True, private_channels=True)
@discord_bot.tree.command(name="model", description="View or switch the current model")
async def model_command(interaction: discord.Interaction, model: str) -> None:
    global curr_model

    if model == curr_model:
        output = f"Current model: `{curr_model}`"
    else:
        if user_is_admin := is_admin_user(interaction.user.id, config):
            curr_model = model
            output = f"Model switched to: `{model}`"
            logging.info(output)
        else:
            output = "You don't have permission to change the model."

    await interaction.response.send_message(output, ephemeral=interaction_is_private(interaction))


@model_command.autocomplete("model")
async def model_autocomplete(interaction: discord.Interaction, curr_str: str) -> list[Choice[str]]:
    global config

    if curr_str == "":
        config = await asyncio.to_thread(get_config)

    choices = [Choice(name=f"â—‰ {curr_model} (current)", value=curr_model)] if curr_str.lower() in curr_model.lower() else []
    choices += [Choice(name=f"â—‹ {model}", value=model) for model in config["models"] if model != curr_model and curr_str.lower() in model.lower()]

    return choices[:25]


@discord.app_commands.allowed_installs(guilds=True, users=True)
@discord.app_commands.allowed_contexts(guilds=True, dms=True, private_channels=True)
@discord_bot.tree.command(name="channelmodel", description="Set this channel's model override")
async def channel_model_command(interaction: discord.Interaction, model: str) -> None:
    global config

    config = await asyncio.to_thread(get_config)

    if not is_admin_user(interaction.user.id, config):
        await interaction.response.send_message("You don't have permission to change channel models.", ephemeral=True)
        return

    if interaction.channel is None:
        await interaction.response.send_message("This command needs a channel.", ephemeral=True)
        return

    if model == "__default__":
        set_channel_override(interaction.channel, "model", None)
        output = f"This channel now uses the global model: `{curr_model}`"
    elif model in config["models"]:
        set_channel_override(interaction.channel, "model", model)
        output = f"This channel's model is now: `{model}`"
    else:
        output = "That model is not in `config.yaml`."

    await interaction.response.send_message(output, ephemeral=interaction_is_private(interaction))


@channel_model_command.autocomplete("model")
async def channel_model_autocomplete(interaction: discord.Interaction, curr_str: str) -> list[Choice[str]]:
    global config, channel_settings

    if curr_str == "":
        config = await asyncio.to_thread(get_config)
        channel_settings = await asyncio.to_thread(load_channel_settings)

    effective_model = get_effective_model(interaction.channel, config)
    choices = []

    if "default" in curr_str.lower() or curr_str == "":
        choices.append(Choice(name=f"Use global default ({curr_model})", value="__default__"))

    if curr_str.lower() in effective_model.lower():
        choices.append(Choice(name=f"Current: {effective_model}", value=effective_model))

    choices += [Choice(name=model, value=model) for model in config["models"] if model != effective_model and curr_str.lower() in model.lower()]
    return choices[:25]


@discord.app_commands.allowed_installs(guilds=True, users=True)
@discord.app_commands.allowed_contexts(guilds=True, dms=True, private_channels=True)
@discord_bot.tree.command(name="persona", description="Set this channel's bot persona")
async def persona_command(interaction: discord.Interaction, persona: str) -> None:
    global config

    config = await asyncio.to_thread(get_config)

    if not is_admin_user(interaction.user.id, config):
        await interaction.response.send_message("You don't have permission to change personas.", ephemeral=True)
        return

    if interaction.channel is None:
        await interaction.response.send_message("This command needs a channel.", ephemeral=True)
        return

    personas = get_personas(config)

    if persona == "__default__":
        set_channel_override(interaction.channel, "persona", None)
        output = "This channel now uses the default persona."
    elif persona in personas:
        set_channel_override(interaction.channel, "persona", persona)
        output = f"This channel's persona is now: `{persona}`"
    else:
        output = "That persona is not in `config.yaml`."

    await interaction.response.send_message(output, ephemeral=interaction_is_private(interaction))


@persona_command.autocomplete("persona")
async def persona_autocomplete(interaction: discord.Interaction, curr_str: str) -> list[Choice[str]]:
    global config, channel_settings

    if curr_str == "":
        config = await asyncio.to_thread(get_config)
        channel_settings = await asyncio.to_thread(load_channel_settings)

    personas = get_personas(config)
    effective_persona = get_effective_persona(interaction.channel, config)
    choices = []

    if "default" in curr_str.lower() or curr_str == "":
        choices.append(Choice(name="Use default persona", value="__default__"))

    if curr_str.lower() in effective_persona.lower():
        choices.append(Choice(name=f"Current: {effective_persona}", value=effective_persona))

    choices += [Choice(name=name, value=name) for name in personas if name != effective_persona and curr_str.lower() in name.lower()]
    return choices[:25]


@discord.app_commands.allowed_installs(guilds=True, users=True)
@discord.app_commands.allowed_contexts(guilds=True, dms=True, private_channels=True)
@discord.app_commands.describe(prompt="What you want the bot to answer", private="Only show the response to you")
@discord_bot.tree.command(name="ask", description="Ask the current model a question")
async def ask_command(interaction: discord.Interaction, prompt: str, private: bool = False) -> None:
    global config

    config = await asyncio.to_thread(get_config)

    if not user_has_permission_for_interaction(interaction, config):
        await interaction.response.send_message("You don't have permission to use this bot here.", ephemeral=True)
        return

    await interaction.response.defer(thinking=True, ephemeral=private)

    try:
        provider_slash_model = get_effective_model(interaction.channel, config)
        persona = get_effective_persona(interaction.channel, config)
        output = await generate_nonstream_response(prompt, interaction.user.id, config, provider_slash_model, persona)
        await send_interaction_chunks(interaction, output, private)
        logging.info(f"/ask completed (user ID: {interaction.user.id}, model: {provider_slash_model}, persona: {persona})")
    except Exception:
        logging.exception("Error while generating /ask response")
        await interaction.followup.send("Something went wrong while generating the response. Check the Render logs.", ephemeral=True)


@discord.app_commands.allowed_installs(guilds=True, users=True)
@discord.app_commands.allowed_contexts(guilds=True, dms=True, private_channels=True)
@discord_bot.tree.command(name="retry", description="Regenerate the bot's most recent channel response")
async def retry_command(interaction: discord.Interaction) -> None:
    global config

    config = await asyncio.to_thread(get_config)

    if not user_has_permission_for_interaction(interaction, config):
        await interaction.response.send_message("You don't have permission to use this bot here.", ephemeral=True)
        return

    if interaction.channel is None:
        await interaction.response.send_message("This command needs a channel.", ephemeral=True)
        return

    await interaction.response.defer(thinking=True, ephemeral=True)

    retry_target = await find_retry_target(interaction.channel)
    if retry_target is None:
        await interaction.followup.send("I couldn't find a previous bot response to retry.", ephemeral=True)
        return

    await interaction.followup.send("Retrying the most recent response in this channel.", ephemeral=True)
    await send_streaming_reply(retry_target, config, log_label="/retry")


@discord.app_commands.allowed_installs(guilds=True, users=True)
@discord.app_commands.allowed_contexts(guilds=True, dms=True, private_channels=True)
@discord.app_commands.describe(message_count="How many recent messages to summarize", private="Only show the summary to you")
@discord_bot.tree.command(name="summarize", description="Summarize recent channel messages")
async def summarize_command(interaction: discord.Interaction, message_count: int = 25, private: bool = False) -> None:
    global config

    config = await asyncio.to_thread(get_config)

    if not user_has_permission_for_interaction(interaction, config):
        await interaction.response.send_message("You don't have permission to use this bot here.", ephemeral=True)
        return

    if interaction.channel is None:
        await interaction.response.send_message("This command needs a channel.", ephemeral=True)
        return

    await interaction.response.defer(thinking=True, ephemeral=private)

    try:
        provider_slash_model = get_effective_model(interaction.channel, config)
        persona = get_effective_persona(interaction.channel, config)
        accept_images = any(x in provider_slash_model.lower() for x in VISION_MODEL_TAGS)
        limited_count = max(1, min(message_count, config.get("max_messages", 25)))
        recent_messages, user_warnings = await build_recent_channel_messages(interaction.channel, config, limited_count, accept_images)

        if not recent_messages:
            await interaction.followup.send("No recent messages found to summarize.", ephemeral=private)
            return

        api_messages = recent_messages[::-1]
        if system_prompt := get_system_prompt(config, persona):
            api_messages.insert(0, dict(role="system", content=system_prompt))
        api_messages.append(dict(role="user", content="Summarize the preceding Discord messages. Include key decisions, open questions, and action items if any. Keep it concise."))

        client_and_kwargs = build_openai_client_and_kwargs(config, provider_slash_model, api_messages, stream=False)
        response = await client_and_kwargs["client"].chat.completions.create(**client_and_kwargs["kwargs"])
        summary = response.choices[0].message.content or "*(empty summary)*"

        if user_warnings:
            summary = "\n".join(sorted(user_warnings)) + "\n\n" + summary

        await send_interaction_chunks(interaction, summary, private)
        logging.info(f"/summarize completed (user ID: {interaction.user.id}, messages: {len(recent_messages)}, model: {provider_slash_model})")
    except Exception:
        logging.exception("Error while generating /summarize response")
        await interaction.followup.send("Something went wrong while generating the summary. Check the Render logs.", ephemeral=True)


@discord_bot.event
async def on_ready() -> None:
    global newspaper_task

    if client_id := config.get("client_id"):
        logging.info(
            f"\n\nBOT SERVER INSTALL URL:\n"
            f"https://discord.com/oauth2/authorize?client_id={client_id}&permissions=412317191168&scope=bot%20applications.commands\n\n"
            f"USER INSTALL URL FOR /ask IN DMS AND GROUP DMS:\n"
            f"https://discord.com/oauth2/authorize?client_id={client_id}&scope=applications.commands&integration_type=1\n"
        )

    if newspaper_task is None or newspaper_task.done():
        newspaper_task = asyncio.create_task(newspaper_scheduler())

    await discord_bot.tree.sync()



def user_has_permission_for_message(new_msg: discord.Message, loaded_config: dict[str, Any]) -> bool:
    is_dm = new_msg.channel.type == discord.ChannelType.private
    role_ids = set(role.id for role in getattr(new_msg.author, "roles", ()))
    channel_ids = set(filter(None, (new_msg.channel.id, getattr(new_msg.channel, "parent_id", None), getattr(new_msg.channel, "category_id", None))))

    permissions = loaded_config["permissions"]
    user_is_admin = is_admin_user(new_msg.author.id, loaded_config)
    if user_is_admin:
        return True

    (allowed_user_ids, blocked_user_ids), (allowed_role_ids, blocked_role_ids), (allowed_channel_ids, blocked_channel_ids) = (
        (perm["allowed_ids"], perm["blocked_ids"]) for perm in (permissions["users"], permissions["roles"], permissions["channels"])
    )

    allow_all_users = not allowed_user_ids if is_dm else not allowed_user_ids and not allowed_role_ids
    is_good_user = user_is_admin or allow_all_users or new_msg.author.id in allowed_user_ids or any(id in allowed_role_ids for id in role_ids)
    is_bad_user = not is_good_user or new_msg.author.id in blocked_user_ids or any(id in blocked_role_ids for id in role_ids)

    allow_dms = loaded_config.get("allow_dms", True)
    allow_all_channels = not allowed_channel_ids
    is_good_channel = allow_dms if is_dm else allow_all_channels or any(id in allowed_channel_ids for id in channel_ids)
    is_bad_channel = not is_good_channel or any(id in blocked_channel_ids for id in channel_ids)

    return not is_bad_user and not is_bad_channel


@discord_bot.event
async def on_message(new_msg: discord.Message) -> None:
    is_dm = new_msg.channel.type == discord.ChannelType.private
    loaded_config = await asyncio.to_thread(get_config)

    if not is_dm and not new_msg.author.bot:
        await track_reputation_message(new_msg, loaded_config)

    if (not is_dm and discord_bot.user not in new_msg.mentions) or new_msg.author.bot:
        return

    if not user_has_permission_for_message(new_msg, loaded_config):
        return

    await send_streaming_reply(new_msg, loaded_config)


@discord_bot.event
async def on_reaction_add(reaction: discord.Reaction, user: discord.User | discord.Member) -> None:
    if user.bot or reaction.message.guild is None or reaction.message.author.bot:
        return
    loaded_config = await asyncio.to_thread(get_config)
    module_config = get_module_config(loaded_config, "reputation_titles", reaction.message.guild.id)
    if not module_config.get("track_reactions", True):
        return
    if not is_public_server_text_channel(reaction.message.channel, set(int(id) for id in module_config.get("ignored_channel_ids", []) or [])):
        return
    await track_reputation_stat(reaction.message.guild.id, reaction.message.author.id, "reactions_received", loaded_config, cooldown_key=("reaction", reaction.message.id, user.id))
    if str(reaction.emoji) in ("🙏", "👍", "💯", "✅"):
        await track_reputation_stat(reaction.message.guild.id, reaction.message.author.id, "thanks_received", loaded_config, cooldown_key=("thanks", reaction.message.id, user.id))


@discord_bot.event
async def on_voice_state_update(member: discord.Member, before: discord.VoiceState, after: discord.VoiceState) -> None:
    if member.bot or member.guild is None or before.channel == after.channel or after.channel is None:
        return
    loaded_config = await asyncio.to_thread(get_config)
    module_config = get_module_config(loaded_config, "reputation_titles", member.guild.id)
    if not module_config.get("track_voice", True):
        return
    await track_reputation_stat(member.guild.id, member.id, "voice_joins", loaded_config, cooldown_key=("voice",))


async def main() -> None:
    await discord_bot.start(config["bot_token"])


try:
    asyncio.run(main())
except KeyboardInterrupt:
    pass
